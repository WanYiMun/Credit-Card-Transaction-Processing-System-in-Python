"""

Introduction to Business Programming : Group Assignment
Group Name : X

Members:
- Alex Tan Yong Xin     2002422  IB
- Pang Tse Ki           2002067  IB
- Wan Yi Mun            2003797  IB

"""


import pandas as pd
import numpy as np
import os
import datetime
import time
import random
import string
import calendar
import docx

#List
loop = True
global_list=[]
choose_list=[]
bank_list  =[]
transaction_file = " "
account_file= " "


#Function


def part_of_day():
    #This function return the time of the day (morning etc.)
    from datetime import datetime
    hour = datetime.now().hour
    if    5 <= hour <= 12:
        return "Good Morning"
    elif 12 <= hour <= 18:
        return "Good Afternoon"
    elif 18 <  hour <= 21:
        return "Good Evening"
    else:
        return "Good Night"

def adminBio():
    #This function return the details of the admin
    passwordDict = {}
    nameDict     = {}
    adminList    = []
    loop1,loop2 = True,True
    with open ("adminBio.txt","r") as f:
        for detail in f: 
            (userId, staff_name, password) = detail.strip("\n").split(" | ")
            passwordDict [userId] = password
            nameDict [userId]     = staff_name
            adminList.append(userId)
    return passwordDict, nameDict, adminList

def adminValidation():
    #This function allow user to login and validated it
    loop1 = True
    passwordDict, nameDict, adminList = adminBio()
    
    while loop1:
        userId = input("User Id  : ")
        if userId.upper() == 'Q':
            quit()
        elif userId not in adminList:
            print ("No such user Id. Please key in the correct user Id")
            loop2 = False
        else:
            password = input("Password : ")
            loop2 = True
            
        while loop2:
            if password.upper() == 'Q':
                break
            
            elif passwordDict[userId] != password:
                print ("Incorrect password. Please key in the correct password")
                password = input("Password : ")
            else:
                loop = True
                return nameDict[userId], loop
                loop1 = False

def login():
    #This function print the login menu and execute the login function
    os.system("cls")
    with open("display2.txt","r") as f:
            content = f.readlines()
            line1   = content[0:5]
            for line in line1:
                print(line.strip("\n"))
    staff_name, loop = adminValidation()
    return staff_name, loop
    print("")

def mainmenu_1():
    #This function print the main menu and allow user to choose their option then validate it
    os.system("cls")
    staff_name, loop = login()
    if loop:
        with open("display2.txt","r") as f:
                content = f.readlines()
                line2   = content[5:14]
                for line in line2:
                    print(line.strip("\n"))
        daytime = part_of_day()
        print ('\n' + daytime, staff_name + ". Have a nice day.\n")
        option = input("Option : ")
        
    while loop:
        if option.upper() == "Q":
            return option, staff_name
        elif option == "1":
            return option, staff_name
        elif option == "2":
            return option, staff_name
        else:
            print("Invalid option. Please key in the correct option")
            option = input("Option : ")

def mainmenu_2():
    #This function print the main menu and allow user to choose their option then validate it
    #This is used when the user enter the main menu for the second time
    os.system("cls")
    global staff_name
    loop = True
    if loop:
        with open("display2.txt","r") as f:
                content = f.readlines()
                line2   = content[5:14]
                for line in line2:
                    print(line.strip("\n"))
        daytime = part_of_day()
        print ('\n' + daytime, staff_name + ". Have a nice day.\n")
        option = input("Option : ")
        
    while loop:
        if option.upper() == "Q":
            loop = False
            print('')
            main()
        elif option == "1":
            return option, loop
        elif option == "2":
            return option, loop
        else:
            print("Invalid option. Please key in the correct option")
            option = input("Option : ")

def cardValidation():
    #This function check whether the customer credit card number exist or not
    global pd, cardNum, file2
    excel_data_df  = pd.read_excel(file2, sheet_name = 'Information', header = 2, usecols = ['Account Id'])
    accountlist    = excel_data_df['Account Id'].tolist()
    if cardNum not in accountlist:
        return "Failed", "Non existing account"
    else:
        return  "Completed", ""

def accValidation():
    #This function check whether the receiver's (transfer target) credit card number exist or not
    global pd, accountID, file2
    excel_data_df  = pd.read_excel(file2, sheet_name = 'Information', header = 2, usecols = ['Account Id'])
    accountlist    = excel_data_df['Account Id'].tolist()
    if accountID not in accountlist:
        return "Failed", "Non existing receiver account"
    else:
        return  "Completed", ""

def statusValidation():
    #This function check whether the status of customer card
    global cardNum, load_workbook, file2
    account_status_dict = {}
    wb  = load_workbook(file2)
    ws = wb['Information']
    rowLength = len(list(ws.rows))
    for i in range (4,rowLength+1):
        accID  = ws.cell(i,1).value
        status = ws.cell(i,10).value
        account_status_dict[accID] = status
    if account_status_dict[cardNum] == "Closed":
        return "Failed", "Account already closed"
    else:
        return "Completed", ""
    
def pMerchIdValid():
    #This function check the merchant id for 001 physical transfer
    global merchantID
    pMerchlist = [1,2,3,4,5,6,7,8]
    if merchantID not in pMerchlist:
        return "Failed","Invalid merchant ID"
    else:
        return "Completed", ""
   
def oMerchIdValid():
    #This function check the merchant id for 004 online transfer
    global merchantID
    oMerchlist = [9,10]
    if merchantID not in oMerchlist:
        return "Failed","Invalid merchant ID"
    else:
        return "Completed", ""

def cLimit_aBalance_1():
    #This function return the credit limit and account balance of customer 
    global cardNum, load_workbook, file2
    account_cLimit_dict = {}
    account_balance_dict = {}
    wb  = load_workbook(file2)
    ws = wb['Information']
    rowLength = len(list(ws.rows))
    for i in range (4,rowLength+1):
        accID  = ws.cell(i,1).value
        cLimit = ws.cell(i,6).value
        balance = ws.cell(i,7).value
        account_cLimit_dict[accID] = cLimit
        account_balance_dict[accID] = balance
    creditlimit = account_cLimit_dict[cardNum]
    balance     = account_balance_dict[cardNum]
    return creditlimit, balance

def cLimit_aBalance_2():
    #This function return the credit limit and account balance of customer, and the account balance receiver (transfer target)
    global cardNum, accountID, load_workbook, file2
    account_cLimit_dict = {}
    account_balance_dict = {}
    wb  = load_workbook(file2)
    ws = wb['Information']
    rowLength = len(list(ws.rows))
    for i in range (4,rowLength+1):
        accID   = ws.cell(i,1).value
        cLimit  = ws.cell(i,6).value
        balance = ws.cell(i,7).value
        account_cLimit_dict[accID] = cLimit
        account_balance_dict[accID] = balance
    creditlimit     = account_cLimit_dict[cardNum]
    accBalance      = account_balance_dict[cardNum]
    receiverBalance = account_balance_dict[accountID]
    return creditlimit, accBalance, receiverBalance

def purchase_withdraw_Validation():
    #This function check whether the transfer amount exceed what the customer can spend or not and return the new balance if does not exist
    global transAmount
    creditLimit, balance = cLimit_aBalance_1()
    new_balance = balance + transAmount
    if new_balance > creditLimit:
        return "Failed","Exceed credit limit", transAmount
    else:
        return "Completed", "", new_balance

def closed_withdraw_Validation():
    #This function check for the customer card which is closed whether they the withdraw amount exceed what the amount customer still have in
    #their account and return the new balance if the transaction is valid
    global transAmount
    creditLimit, balance = cLimit_aBalance_1()
    if balance < 0:
        new_balance = balance + transAmount
        if new_balance <= 0:
            return "Completed", "", new_balance
        else:
            return "Failed","Exceed account balance", transAmount
    else:
        return "Failed", "Account already closed", transAmount

def paymentValidation():
    #This function return the new balance when customer pay their credit card bill
    global transAmount
    creditLimit, balance = cLimit_aBalance_1()
    new_balance = balance - transAmount
    return "Completed", "", new_balance

def transferValidation():
    #This function validate the transfer process 003 and return the customer and receiver new balance 
    global transAmount
    creditLimit, accBalance, receiverBalance = cLimit_aBalance_2()
    new_acc_balance = accBalance + transAmount
    if new_acc_balance > creditLimit:
        return "Failed","Exceed credit limit", "", ""
    else:
        new_receiver_balance = receiverBalance - transAmount
        return "Completed", "", new_acc_balance, new_receiver_balance
 
def current_pin():
    #This function return the current pin of customer 
    global load_workbook, file2, cardNum
    wb  = load_workbook(file2)
    ws  = wb['Information']
    current_pin_dict = {}
    rowLength = len(list(ws.rows))
    for i in range (4,rowLength+1):
        accID  = ws.cell(i,1).value
        current_pin = ws.cell(i,4).value
        current_pin_dict[accID] = current_pin
    current_pin = current_pin_dict[cardNum]
    return current_pin
    
def update_pin():
    #This function update the new pin and old pin of the customer in GA2_Accounts.xlsx
    global file2, load_workbook, pd,cardNum, new_pin, old_pin
    excel_data_df  = pd.read_excel(file2, sheet_name = 'Information', header = 2, usecols = ['Account Id'])
    wb  = load_workbook(file2)
    ws  = wb['Information']
    accountlist              = excel_data_df['Account Id'].tolist()
    row_index                = accountlist.index(cardNum)+4
    current_pin_position     = 'D' + str(row_index)
    old_pin_position         = 'I' + str(row_index)
    ws[current_pin_position] = new_pin
    ws[old_pin_position]     = old_pin
    wb.save(file2)

def current_address():
    #This function return the current address of the customer 
    global load_workbook, file2, cardNum
    wb  = load_workbook(file2)
    ws  = wb['Information']
    current_add_dict = {}
    rowLength = len(list(ws.rows))
    for i in range (4,rowLength+1):
        accID  = ws.cell(i,1).value
        current_add = ws.cell(i,3).value
        current_add_dict[accID] = current_add
    current_address = current_add_dict[cardNum]
    return current_address
    
def update_address():
    #This function update the new address and old address of the customer in GA2_Accounts.xlsx
    global file2, load_workbook, pd,cardNum, new_address, old_address
    excel_data_df  = pd.read_excel(file2, sheet_name = 'Information', header = 2, usecols = ['Account Id'])
    wb  = load_workbook(file2)
    ws  = wb['Information']
    accountlist              = excel_data_df['Account Id'].tolist()
    row_index                = accountlist.index(cardNum)+4
    current_address_position     = 'C' + str(row_index)
    old_address_position         = 'H' + str(row_index)
    ws[current_address_position] = new_address
    ws[old_address_position]     = old_address
    wb.save(file2)

def update_balance_1():
    #This function update the new balance of the customer for trans type 001, 002, 003(payment) in GA2_Accounts.xlsx
    global file2,load_workbok, pd, cardNum, new_balance
    excel_data_df  = pd.read_excel(file2, sheet_name = 'Information', header = 2, usecols = ['Account Id'])
    wb  = load_workbook(file2)
    ws  = wb['Information']
    accountlist  = excel_data_df['Account Id'].tolist()
    row_index    = accountlist.index(cardNum)+4
    position     = 'G' + str(row_index)
    ws[position] = new_balance
    wb.save(file2)

def update_balance_2():
    #This function update the new balance of the customer and receiver for trans type 003(transfer) in GA2_Accounts.xlsx
    global file2, load_workbook, pd, cardNum, accountID, new_acc_balance, new_receiver_balance
    excel_data_df  = pd.read_excel(file2, sheet_name = 'Information', header = 2, usecols = ['Account Id'])
    wb  = load_workbook(file2)
    ws  = wb['Information']
    accountlist    = excel_data_df['Account Id'].tolist()
    row_index_1    = accountlist.index(cardNum)+4
    position_1     = 'G' + str(row_index_1)
    row_index_2    = accountlist.index(accountID)+4
    position_2     = 'G' + str(row_index_2)
    ws[position_1] = new_acc_balance
    ws[position_2] = new_receiver_balance
    wb.save(file2)

def update_status_remark():
    #This function update the status and remark of the transaction in the excel which the admin key in to proccess 
    global index, status, remark, file1, sheet
    wb  = load_workbook(file1)
    ws  = wb[sheet]
    status_position = 'L' + str(index+3)
    remark_position = 'M' + str(index+3)   
    ws[status_position] = status
    ws[remark_position] = remark
    wb.save(file1)

def open_account_validation():
    #This is the valid the proccess 008
    status, remark = cardValidation()
    if status == 'Failed':
        status, remark = 'Completed', ''
        return status, remark
    else:
        status, remark = 'Failed', ''
        return status, remark

def update_new_account():
    #This is to create a new record of new customer in GA2_Accounts.xlsx
    global file2, load_workbook,pd, new_account_data
    wb   = load_workbook(file2)
    ws = wb['Information']
    ws.append(new_account_data)
    wb.save(file2)

def update_open_account():
    #This is to change account status from closed to active and update their credit limit
    global file2,load_workbok, pd, cardNum, new_balance, transAmount
    excel_data_df  = pd.read_excel(file2, sheet_name = 'Information', header = 2, usecols = ['Account Id'])
    wb  = load_workbook(file2)
    ws  = wb['Information']
    accountlist  = excel_data_df['Account Id'].tolist()
    row_index    = accountlist.index(cardNum)+4
    status_position     = 'J' + str(row_index)
    credit_position     = 'F' + str(row_index)
    ws[status_position] = 'Active'
    ws[credit_position] = transAmount
    wb.save(file2)

def update_close_account():
    #This is to change the accoutn status from active to closed
    global file2,load_workbok, pd, cardNum, new_balance
    excel_data_df  = pd.read_excel(file2, sheet_name = 'Information', header = 2, usecols = ['Account Id'])
    wb  = load_workbook(file2)
    ws  = wb['Information']
    accountlist  = excel_data_df['Account Id'].tolist()
    row_index    = accountlist.index(cardNum)+4
    position     = 'J' + str(row_index)
    ws[position] = 'Closed'
    wb.save(file2)
    
def year():
    #This is to get and return the year of each transaction
    global file1, sheet, pd
    excel_data_df  = pd.read_excel(file1, sheet_name = sheet, header = 2)
    dateFormat = excel_data_df ['Date (dd/mm/yyy)'].apply(lambda x: pd.Timestamp(x).strftime('%Y-%m-%d'))
    yearList   = []
    for row in range(len(excel_data_df)):
        yymmdd   = dateFormat[row]
        tempList = yymmdd.split('-')
        year     = tempList[0]
        yearList.append(year)
    return yearList

def date_time():
    #This is to get the date and time of the day
    from datetime import datetime
    date_time = datetime.now().strftime('%d/%m/%y %H:%M:%S %p')
    return date_time

def update_excel():
    #This is to append all the processed transaction into the main excel file Processed Transaction.xlsx based on the year of transaction
    global file3, load_workbook, processed_transaction
    wb   = load_workbook(file3)
    ws1  = wb['2019']
    ws2  = wb['2020']
    ws3  = wb['2021']
    global processed_transaction
    yearList = year()

    for i in range(len(yearList)):
        if yearList[i] == '2019':
            ws1.append(processed_transaction[i])
        elif yearList[i] =='2020':
            ws2.append(processed_transaction[i])
        elif yearList[i] == '2021':
            ws3.append(processed_transaction[i])
    wb.save(file3)

def option1():
    #This is to execute the functions for transfer validatio and update the excel file to process the transaction
    global pd, file1, sheet, date, time, cardNum, transactionType, transAmount, merchantID, accountID,new_address, new_pin,name,nric, status, remark
    global new_balance, processed_transaction,new_receiver_balance, new_acc_balance, index,old_pin, old_address,new_account_data
    import pandas as pd
    excel_data_df  = pd.read_excel(file1, sheet_name = sheet, header = 2)
    dateList       = excel_data_df['Date (dd/mm/yyy)'].tolist()
    timeList       = excel_data_df['Time'].tolist()
    cardList       = excel_data_df['Credit Card Number'].tolist()
    tTypeList      = excel_data_df['Tran.Type'].tolist()
    tAmountList    = excel_data_df['Tran.Amount/Credit Limit'].tolist()
    merchIdList    = excel_data_df['Merchant Id number'].tolist()
    accIDList      = excel_data_df['Account Id Number'].tolist()
    newPinList     = excel_data_df['New Pin'].tolist()
    newAddList     = excel_data_df['New Address'].tolist()
    nameList       = excel_data_df['Name'].tolist()
    nricList       = excel_data_df['IC. No'].tolist()
    statusList     = excel_data_df['Status'].tolist()
    remarkList     = excel_data_df['Remarks'].tolist()
    data_row_list  = []
    processed_transaction = []
    index = 0
    
    for row in range(len(excel_data_df)):
        data_row_list.append([dateList[row],timeList[row],cardList[row],tTypeList[row],tAmountList[row],merchIdList[row],accIDList[row],newAddList[row],newPinList[row],nameList[row],nricList[row],statusList[row],remarkList[row]])


    for t in data_row_list:
        date            = t[0]
        time            = t[1]
        cardNum         = t[2]
        transactionType = t[3]
        transAmount     = t[4]
        merchantID      = t[5]
        accountID       = t[6]
        new_address     = t[7]
        new_pin         = t[8]
        name            = t[9]
        nric            = t[10]
        status          = t[11]
        remark          = t[12]
        index           += 1
        
        if status == 'Pending':
            status,remark   = cardValidation()
        
            if status == 'Completed':
                status,remark   = statusValidation()

            if status == 'Completed' and transactionType == 1:
                status,remark   = pMerchIdValid()
                if status == 'Completed':
                    status, remark, new_balance   = purchase_withdraw_Validation()
                    update_status_remark()
                    if status == 'Completed':
                        remark = date_time()
                        update_status_remark()
                        update_balance_1()
                else:
                    update_status_remark()
                    
            elif (status == 'Completed' or 'Failed') and transactionType == 2:
                if status == 'Completed':
                    status, remark, new_balance   = purchase_withdraw_Validation()
                    update_status_remark()
                    if status == 'Completed':
                        remark = date_time()
                        update_status_remark()
                        update_balance_1()
                else:
                    status, remark, new_balance   = closed_withdraw_Validation()
                    update_status_remark()
                    if status == 'Completed':
                        remark = date_time()
                        update_status_remark()
                        update_balance_1()
                          
            elif (status == 'Completed' or 'Failed')  and transactionType == 3:
                if cardNum == accountID:
                    status, remark, new_balance = paymentValidation()
                    remark = date_time() 
                    update_status_remark()
                    update_balance_1()
                else:
                    if status == 'Completed':
                        status, remark = accValidation()
                        update_status_remark()
                        if status == 'Completed':
                            status, remark, new_acc_balance, new_receiver_balance = transferValidation()
                            update_status_remark()
                            if status == 'Completed':
                                remark = date_time()
                                update_status_remark()
                                update_balance_2()
                    else:
                        status, remark = "Failed","Account already closed"
                        update_status_remark()
                        
                        
            elif status == 'Completed' and transactionType == 4:
                status,remark   = oMerchIdValid()
                if status == 'Completed':
                    status, remark, new_balance   = purchase_withdraw_Validation()
                    update_status_remark()
                    if status == 'Completed':
                        remark = date_time()
                        update_status_remark()
                        update_balance_1()
                else:
                    update_status_remark()
                    
            elif status == 'Completed' and transactionType == 6:
                old_address = current_address()
                if new_address != old_address:
                    remark = date_time()
                    update_status_remark()
                    update_address()
                else:
                    status,remark = 'Failed', 'New address same with old address'
                    update_status_remark()

            elif status == 'Completed' and transactionType == 7:
                old_pin = current_pin()
                if int(new_pin) != int(old_pin):
                    remark = date_time()
                    update_status_remark()
                    update_pin()
                else:
                    status,remark = 'Failed', 'New pin same with old pin'
                    update_status_remark()
                    
            elif (status == 'Completed' or 'Failed') and transactionType == 8:
                status, remark = open_account_validation()
                if status == 'Completed':
                    account_balance = 0
                    old_add = ''
                    old_pin = ''
                    card_status = 'Active'
                    new_account_data = [cardNum,name,new_address,new_pin,nric,transAmount,account_balance,old_add,old_pin,card_status]
                    status,remark = 'Completed',date_time()
                    update_status_remark()
                    update_new_account()
                else:
                    status,remark = statusValidation()
                    if status == 'Failed':
                        status,remark = 'Completed',''
                        remark = date_time()
                        update_status_remark()
                        update_open_account()
                    else:
                        status,remark = 'Failed','Account already active'
                        update_status_remark()
                    
            elif (status == 'Completed' or 'Failed')  and transactionType == 9:
                status,remark = statusValidation()
                if status == 'Completed':
                    status,remark = 'Completed',date_time()
                    update_status_remark()
                    update_close_account()
                else:
                    status,remark = 'Failed','Account already closed'
                    update_status_remark()

            else:
                if (status == 'Completed') and (transactionType not in [1,2,3,4,5,6,7,8,9]):
                    status, remark  = "Failed","Invalid transaction type"
                    update_status_remark()
                elif (status == 'Failed') and (transactionType in [1,4,5,6,7]):
                    update_status_remark()
                elif (status == 'Failed') and (transactionType not in [2,3,8,9]):
                    status, remark  = "Failed","Invalid transaction type"
                    update_status_remark()   
            
        processed_transaction.append([date,time,cardNum,transactionType,transAmount,merchantID,accountID,new_address,new_pin,name,nric,status,remark])
    update_excel()   



# [ Read Updateted Excel and transfer to List ]
##Those Function is to read the Excel file and turn the record into list . 
def read_updated_transaction_file(transaction_file):
    df= pd.concat(pd.read_excel(transaction_file, sheet_name=None,header = 2), ignore_index=True)    
    df.fillna("Null", inplace=True)
    df['Date (dd/mm/yyy)'] = pd.to_datetime(df['Date (dd/mm/yyy)']) 
    df['Year'] = df['Date (dd/mm/yyy)'].dt.year
    df['Month'] = df['Date (dd/mm/yyy)'].dt.month
    tran_year_list       = df['Year'].tolist() 
    tran_month_list      = df['Month'].tolist()
    tran_type_list       = df['Tran.Type'].tolist() 
    tran_merchant_id_list= df['Merchant Id number'].tolist()
    tran_account_id_list = df['Account Id Number'].tolist()
    tran_credit_card_list = df['Credit Card Number'].tolist()
    tran_status_list     = df['Status'].tolist()
    tran_amount_list     = df['Tran.Amount/Credit Limit'].tolist()
    return tran_year_list,tran_month_list ,tran_type_list,tran_merchant_id_list,tran_account_id_list,tran_status_list,tran_amount_list ,tran_credit_card_list

def read_updated_account_file(account_file):
    df = pd.concat(pd.read_excel(account_file, sheet_name=None,header = 2), ignore_index=True)
    acc_id_list          = df['Account Id'].tolist()
    acc_name_list        = df['Name'].tolist()
    acc_ic_list          = df['IC.No'].tolist()
    acc_credit_list      = df['Credit Limit'].tolist()
    acc_balance_list     = df['Account Balance(Amount)'].tolist()
    acc_status_list      = df['Credit .C Status'].tolist()
    acc_address_list     = df['Current Address'].tolist()

    return acc_id_list,acc_name_list,acc_ic_list,acc_credit_list,acc_balance_list,acc_status_list,acc_address_list


# [Count Page (Credit Card Transaction Record )]
def count_page():
    count_correct=0
    count_error=0
    tran_year_list,tran_month_list ,tran_type_list,tran_merchant_id_list,tran_account_id_list,tran_status_list,tran_amount_list ,tran_credit_card_list = read_updated_transaction_file(transaction_file)
    if (len(tran_status_list)!=0) and (len(tran_status_list)>0) :
        for i in range (len(tran_status_list)):
            if (tran_status_list[i]== "Completed") and (tran_credit_card_list[i][2] == bank_list[0])  :
                count_correct+=1
            elif (tran_status_list[i]== "Failed") and (tran_credit_card_list[i][2] == bank_list[0]):
                count_error+=1
        b1=(count_error/20)
        b2=(count_error//20)
        if b1>b2:
            error_pages=b2+1
        elif b1 == b2 :
            error_pages=b2
        elif b2 == 0 or b1== 0 :
            error_pages = 1
        a1=(count_correct/20)
        a2=(count_correct//20)
        if a1>a2:
            correct_pages=a2+1
        elif a1 == a2 :
            correct_pages=a2
        elif a2 == 0 or a1 ==0 :
            correct_pages = 1
    elif (len(tran_status_list)==0):
        correct_pages = 1
        error_pages =1
    return correct_pages ,error_pages

# [Count Page (Credit Card Account Information)]
def count_acc_page():
    acc_id_list,acc_name_list,acc_ic_list,acc_credit_list,acc_balance_list,acc_status_list,acc_address_list = read_updated_account_file(account_file)
    count_correct=0
    count_error=0
    if (len(acc_status_list) != 0 and (len(acc_status_list) >10)):
        for i in range (len(acc_status_list)):
            if (acc_status_list[i]== "Active" ) and (acc_id_list[i][2] == bank_list[0]) :
                count_correct+=1
            elif (acc_status_list[i]== "Closed" ) and (acc_id_list[i][2] == bank_list[0]):
                count_error+=1
        b1=(count_error/20)
        b2=(count_error//20)
        if b1>b2:
            error_pages=b2+1
        elif b1 == b2 :
            error_pages=b2
        elif b2 == 0 or b1== 0 :
            error_pages = 1
        a1=(count_correct/20)
        a2=(count_correct//20)
        if a1>a2:
            correct_pages=a2+1
        elif a1 == a2 :
            correct_pages=a2
        elif a2 == 0 or a1 ==0 :
            correct_pages = 1
    elif (len(acc_status_list) == 0):
        correct_pages = 1
        error_pages =1
    return correct_pages ,error_pages


#1. [ Bank Option ] - Public Bank , MayBank , RHB
#Due to the credit card transaction record may have account from public bank (001-XXXX) , maybank (002-XXXX) or RHB (003-XXXX) ,
# this function allows the user to select which Bank's credit card transaction record that user wants to view .
def bank_option():
    os.system("cls")
    option_list=["1","2","3" ,"4"]
    with open ("display.txt" , "r") as f :
        for line in f.readlines()[0:8]:
            print(line)
    option=input("Option = ")
    while (option not in option_list):
        print ("\t[Please Try Again]") 
        option=input("Option = ")
    if (option == "1") or (option == "2" ) or (option == "3"):
        bank_list.append(str(option)) 
        choose_record()
    elif (option == "4"):
        bank_list.clear()
        main()


# 2. [ Choose Report Type ] - Credit Card Transaction Record // Account Information and Balance
def choose_record():
    os.system("cls")
    option_list=["1","2","3"]
    with open ("display.txt" , "r") as f :
        for line in f.readlines()[7:14]:
            print(line)
    option=input("Option = ")
    while (option not in option_list):
        print ("\t[Please Try again.]") 
        option=input("Option = ")
    if (option == "1"):
        report_option()
    if (option == "2" ):
        report1_option()
    elif (option == "3"):
        bank_list.clear()
        bank_option()


# 3. [Choose Credit Card Transaction Record ] - No Error Report // Error Report
def report_option():
    os.system("cls")
    option_list=["1","2","3"]
    with open ("display.txt" , "r") as f :
        for line in f.readlines()[13:20]:
            print(line)
    option=input("Option = ")
    while (option not in option_list):
        print ("\t[Please Try Again ]") 
        option=input("Option = ")
    if (option == option_list[0]):
        correct_pages ,error_pages = count_page()
        tran_df ,tran_df2=transaction_report(correct_pages,transaction_file)
    elif (option == option_list[1]):
        correct_pages ,error_pages = count_page() 
        tran_error_df ,tran_error_df2=transaction_error_report(error_pages,transaction_file)
    elif (option == option_list[2]):
        choose_record() 



## A.1 [ Transaction Report ]  - No Error Record
# This function allows user to view all the transaction record  .
def transaction_report(c_pages,tran_file):
    os.system("cls")
    global calendar
    group_n=[]
    tran_year_list,tran_month_list ,tran_type_list,tran_merchant_id_list,tran_account_id_list,tran_status_list,tran_amount_list ,tran_credit_card_list = read_updated_transaction_file(transaction_file)
    cols=[0,2,3,4,5,6,11,12]
    tran_df= pd.concat(pd.read_excel(tran_file, sheet_name=None,header = 2 ,usecols=cols), ignore_index=True)
    tran_df['Merchant Id number'] = tran_df['Merchant Id number'].fillna(0).astype(int)
    tran_df['Merchant Id number'] = (tran_df['Merchant Id number'].astype(str))
    tran_df['Merchant Id number'] = (tran_df['Merchant Id number'].str.rjust(3, "0"))
    tran_df['Merchant Id number']  = tran_df['Merchant Id number'].replace("000",'Null')            
    tran_df['Tran.Type'] = tran_df['Tran.Type'].fillna(0).astype(int)
    tran_df['Tran.Type'] = (tran_df['Tran.Type'].astype(str))
    tran_df['Tran.Type'] = (tran_df['Tran.Type'].str.rjust(3, "0"))
    tran_df['Tran.Type']  = tran_df['Tran.Type'].replace("000",'Null')
    tran_df['Tran.Amount/Credit Limit'] = tran_df['Tran.Amount/Credit Limit'].fillna(0).astype(float)
    tran_df['Tran.Amount/Credit Limit'] = ('RM' + tran_df['Tran.Amount/Credit Limit'].astype(str))
    tran_df['Date (dd/mm/yyy)'] = pd.to_datetime(tran_df['Date (dd/mm/yyy)'])
    tran_df.fillna("Null", inplace=True)
    pd.set_option('display.max_rows', None)
    pd.set_option('display.max_columns', None)
    pd.set_option('display.width', None)
    pd.set_option('display.max_colwidth', None)
    if (bank_list[0] == "1"):
        tran_df = tran_df[tran_df['Credit Card Number'].str.startswith("001", na = False)]
    elif(bank_list[0] == "2"):
        tran_df = tran_df[tran_df['Credit Card Number'].str.startswith("002", na = False)]
    elif (bank_list [0] == "3"):
        tran_df = tran_df[tran_df['Credit Card Number'].str.startswith("003", na = False)]
    tran_df2=tran_df[tran_df["Status"]== "Completed" ]
    if (len(tran_df2)==0):
        print("\n")
        print("                                           ABC Company                                 ", datetime.date.today() ,time.strftime("%H:%M:%S"))
        print("-"*120)
        print("Report   : Transactions Report  (ALL) ")
        print("Address  :  122,Jalan Ikan  , 20111 Kuala Kangsa , Malaysia")
        print("Year     :",(min(tran_year_list)),"to",(max(tran_year_list)))
        print("Month    :",calendar.month_name[(min(tran_month_list))],"to",calendar.month_name[(max(tran_month_list))],"   "*21,"[ Page = 1]")
        print("-"*120)
        print("\n\n")
        print("                    [Empty Record]            ")
        print(" "*80,"(Q)uit")
        print("-"*120)
        option = input("Option = :").upper()
        while (option != "Q"):
            print ("\t[Please enter again.]")
            option=input("Option = ").upper()
        if option == "Q" :
            choose_list.clear()
            report_option()
    else:
        n= (len(tran_df2)//c_pages)
        group_n.append(0)
        if (isinstance(n,int)) == True :
            for i in range (c_pages):
                group_n.append(n)
        else:
            for i in range (c_pages):
                if (c_pages - i) != 1:
                    n= (len(tran_df2)//c_pages)
                    group_n.append(n)
                else:
                    special_n=(len(tran_df2)-(i*n))
                    group_n.append(special_n)
        groups = np.cumsum(group_n)
        d = {}
        for i in range(c_pages):
            print("\n")
            print("                                           ABC Company                                 ", datetime.date.today() ,time.strftime("%H:%M:%S"))
            print("-"*120)
            print("Report   : Transactions Report  (ALL) ")
            print("Address  :  122,Jalan Ikan  , 20111 Kuala Kangsa , Malaysia")
            print("Year     :",(min(tran_year_list)),"to",(max(tran_year_list)))
            print("Month    :",calendar.month_name[(min(tran_month_list))],"to",calendar.month_name[(max(tran_month_list))],"   "*21,"[ Page",c_pages,"of",i+1,"]")
            print("-"*120)
            print("\n\n")
            d[i+1] = tran_df2.iloc[groups[i]:groups[i+1]]
            print(d[i+1])
            print("\n","  "*60,"Total record =",(len(tran_df2)),"\n")
            print(" "*70,"         (N)ext   (S)ort   (P)rint   (Q)uit")
            print("-"*120)
            option=input("Option = ").upper()
            option_list=["N","S","P","Q"]
            while option not in option_list:
                print ("\t[Please enter again.]")
                option=input("Option = ").upper()
            if(option == "N"):
                if (c_pages - i) != 1:
                    os.system("cls")
                    continue
                else:
                    transaction_report(c_pages,tran_file)
            elif (option == "S"):
                group_n.clear()
                sort_option_list=["1","2","3"]
                with open ("display.txt" , "r") as f :
                    for line in f.readlines()[19:26]:
                        print(line)
                sort_option=input("Option = ")
                while (sort_option not in sort_option_list) :
                    print ("\t[Please enter again.]")
                    sort_option=input("Option = ")
                if (sort_option == "1"):
                    tran_sort_option()
                elif (sort_option == "2"):
                    tran_sort_option1()
                elif (sort_option == "3"):
                    transaction_report(c_pages,tran_file)
            elif (option == "P"):
                group_n.clear()
                tran_print_option(tran_df,tran_df2,c_pages)
            elif (option == "Q"):
                group_n.clear()
                report_option()
    return tran_df,tran_df2



### A.1.1 - Print the Credit Card Transaction Record in TEXT file // WORD file // PDF file (No Error)
###       - Transaction Type = ALL
def tran_print_option(data1,data,c_pages):
    os.system("cls")
    correct_pages ,error_pages = count_page()
    group_n=[]
    tran_year_list,tran_month_list ,tran_type_list,tran_merchant_id_list,tran_account_id_list,tran_status_list,tran_amount_list ,tran_credit_card_list = read_updated_transaction_file(transaction_file)
    with open ("display.txt" , "r") as f :
        for line in f.readlines()[25:33]:
            print(line)
    option=input("Option = ")
    filename="Transaction File_"+""+str(time.strftime("%Y%m%d-%H%M%S"))
    option_list=["1","2","3","4"]
    while option not in option_list:
        print ("\t[Please enter again.]")
        option=input("Option = ")
    if (option =="1"):
        filename=filename + ".txt"        
        n= (len(data)//c_pages)
        group_n.append(0)
        if (isinstance(n,int)) == True :
            for i in range (c_pages):
                group_n.append(n)
        else:
            for i in range (c_pages):
                if (c_pages - i) != 1:
                    n= (len(data)//c_pages)
                    group_n.append(n)
                else:
                    special_n=(len(data)-(i*n))
                    group_n.append(special_n)
        groups = np.cumsum(group_n)
        with open (filename,"w") as f :
            for i in range(c_pages):
                d = {}
                d[i+1]= data.iloc[groups[i]:groups[i+1]]
                f.write("\n\n\n")
                f.write("                                           ABC Company                                         ")
                f.write(str(datetime.date.today()))
                f.write(str(time.strftime("%H:%M:%S")))
                f.write("\n")
                f.write("-"*120)
                f.write("\n")
                f.write("Report   : Transactions Report (ALL) \n")
                f.write("Transaction Type : All")
                f.write("Address  :  122,Jalan Ikan , 20111 Kuala Kangsa , Malaysia\n")
                f.write("Reported by  : Alex Pan Wan ")
                f.write("Year     : ")
                f.write(str((min(tran_year_list))))
                f.write(" to ")
                f.write(str((max(tran_year_list))))
                f.write("\n")
                f.write("Month    : ")
                f.write(str(calendar.month_name[(min(tran_month_list))]))
                f.write(" to ")
                f.write(str(calendar.month_name[(max(tran_month_list))]))
                f.write("   "*21)
                f.write("Page ")
                f.write(str(c_pages))
                f.write(" of ")
                f.write(str(i+1))
                f.write("\n")
                f.write("-"*120)
                f.write("\n")
                f.write(str(d[i+1]))
                f.write("\n\n")
                f.write("  "*60)
                f.write("Total record = ")
                f.write(str(len(data)))
                f.write("\n")
                f.write("-"*120)
                f.write("\n\n\n")
        print("Please Check your file >>> File Name = ",filename)
        print("Location : Documents // this File \n\n .......Thank You.....")
        enter=input("Press Enter to continue :)")
        group_n.clear()
        tran_print_option(data1,data,c_pages)
    elif (option =="2") or (option == "3"):
        print("...................Processing...................")
        print(" .................Please Wait..................")
        from docx import Document
        from docx.shared import Inches
        word_filename= filename + ".docx"
        document = Document()
        document.add_heading("                        [ABC Company]         ",0).bold = True
        document.add_paragraph("Report   : Transactions Report (ALL) ").bold=True
        document.add_paragraph("Address  :  122,Jalan Ikan , 20111 Kuala Kangsa , Malaysia")
        document.add_paragraph("Reported by : Alex Pan Wan")
        document.add_paragraph("_________________________________________________________________________________________________________")
        document.add_paragraph(str(datetime.date.today()))
        document.add_paragraph(str(time.strftime("%H:%M:%S")))
        t = document.add_table(data.shape[0]+1, data.shape[1])
        for j in range(data.shape[-1]):
            t.cell(0,j).text = data.columns[j]
        for i in range (data.shape[0]):
            for j in range(data.shape[-1]):
                t.cell(i+1,j).text = str (data.values[i,j])
        document.save(str(word_filename))
        if (option == "2"):
            print("Please Check your file >>> File Name = ",word_filename)
            enter=input("Press Enter to continue :)")
            group_n.clear()
            tran_print_option(data1,data,c_pages)
        elif (option == "3"):
            from docx2pdf import convert
            pdf_filename=filename + ".pdf"
            convert(word_filename,pdf_filename)
            os.remove(word_filename)
            print("Please Check your file >>> File Name = ",pdf_filename)
            print("Location : Documents // this File \n\n............ Thank You.............")
            enter=input("Press Enter to continue :)")
            group_n.clear()
            tran_print_option(data1,data,c_pages)
    elif (option == "4"):
        group_n.clear()
        transaction_report(correct_pages,transaction_file)



### A.1.2 - Sort Option = 1. Transaction Type + Month + Year
##This function allows user to select transaction type + month + year .
def tran_sort_option( ):
    os.system("cls")
    correct_pages ,error_pages = count_page()
    option_list=["1","2","3","4","6","7","8","9","10","11"]
    tran_year_list,tran_month_list ,tran_type_list,tran_merchant_id_list,tran_account_id_list,tran_status_list,tran_amount_list ,tran_credit_card_list = read_updated_transaction_file(transaction_file)
    group_n=[]
    global choose_list
    global global_list
    global loop
    with open ("display.txt" , "r") as f :
        for line in f.readlines()[32:46]:
            print(line)
    option=input("Option = ")
    while (option not in option_list):
        print ("\t[Please enter again.]")
        option=input("Option = ")
    if (option != "10")and (option !="11"):
        t_type = "00" + option
        choose_list.append(t_type)
    if (option == "10"):
        t_type= "0" + option
        choose_list.append(t_type)
    if (option != "11"):
        print("-------------------------------------------------------------------")
        print("                        Choose Year")
        print("-------------------------------------------------------------------")
        space1_list = tran_year_list.copy()
        space1_list.sort()
        for i in range (len(space1_list)):
            if (space1_list [i] not in global_list):
                global_list.append(space1_list[i])
                print("                 ","*",space1_list[i])
        print("-------------------------------------------------------------------")
        while loop :
            try:
                year=input("Year (eg:20XX) = ")
                if (int(year) in tran_year_list):
                    loop = False
                else:
                    loop = True
                    print ("\t[Please enter agin.]")
            except ValueError:
                loop = True
                print ("\t[Please enter again.]")
        choose_list.append(int(year))
        global_list.clear()
        loop = True 
        print("-------------------------------------------------------------------")
        print("                        Choose Month")
        print("-------------------------------------------------------------------")
        space2_list = tran_month_list.copy()
        space2_list.sort()
        for i in range (len(space2_list)):
            if (space2_list [i] not in global_list):
                global_list.append(space2_list[i])
                print("                 *",space2_list[i],"=",calendar.month_name[space2_list[i]])
        print("-------------------------------------------------------------------")
        while loop :
            try:
                month=input("Month (eg:1) = ")
                if (int(month) in tran_month_list):
                    loop = False
                else:
                    loop = True
                    print ("\t[Please enter again.]")
            except ValueError:
                loop = True
                print ("\t[Please enter again.]")
        loop = True
        choose_list.append(int(month))
        global_list.clear()
        if (choose_list[0] == "001") or (choose_list[0]=="004"):
            cols=[0,2,3,4,5,11,12]
            tran_sort_option_display(transaction_file,cols)
        elif (choose_list[0] == "002" ):
            cols=[0,2,3,4,11,12]
            tran_sort_option_display(transaction_file,cols)
        elif (choose_list[0] == "003"):
            cols=[0,2,3,4,6,11,12]
            tran_sort_option_display(transaction_file,cols)
        elif (choose_list[0] == "007" ) or (choose_list[0] == "009") or (choose_list[0]=="006"):
            cols=[0,2,3,11,12]
            tran_sort_option_display(transaction_file,cols)
        elif (choose_list[0] == "008"):
            cols=[0,2,3,4,5,11,12]
            tran_sort_option_display(transaction_file,cols)
        elif (choose_list[0] == "010"):
            cols=[0,2,3,4,5,6,11,12]
            tran_sort_option_display_10(transaction_file,cols)
    else:
        group_n.clear()
        choose_list.clear()
        transaction_report(correct_pages,transaction_file)



### A.1.2.1 - Print Sorted Credit Card Transaction Record + Create text file // PDF file // WORD file (No Error)
###         - Transaction Type = 1,2,3,4,6,7,8,9
## This function will display all the transaction record that user choose based on transaction type + month + year and user can print the sorted transaction record . 
def tran_sort_option_display(tran_file,column):
    os.system("cls")
    global choose_list
    group_n =[]
    tran_year_list,tran_month_list ,tran_type_list,tran_merchant_id_list,tran_account_id_list,tran_status_list,tran_amount_list ,tran_credit_card_list = read_updated_transaction_file(transaction_file)
    correct_pages ,error_pages = count_page()
    transaction_type={"1": "Purchase (Physical)" , "2" : "Cash Withdrawal", "3": "Payment/Deposit into Account","4":"Purchase (Online)"  , "6" : "change of address", "7" : "change PIN numbers" , "8" : "Open/Approval of New Account" , "9": "Close an Account" }
    tran_sort_df = pd.concat(pd.read_excel(tran_file, sheet_name=None,header = 2 ,usecols=column), ignore_index=True)
    if (choose_list[0][2] == "8") or (choose_list[0][2] == "1") or (choose_list[0][2]=="4"):
        tran_sort_df['Merchant Id number'] = tran_sort_df['Merchant Id number'].fillna(0).astype(int)
        tran_sort_df['Merchant Id number'] = (tran_sort_df['Merchant Id number'].astype(str))
        tran_sort_df['Merchant Id number'] = (tran_sort_df['Merchant Id number'].str.rjust(3, "0"))
        tran_sort_df['Merchant Id number']  = tran_sort_df['Merchant Id number'].replace("000",'Null')
        tran_sort_df['Tran.Amount/Credit Limit'] =tran_sort_df['Tran.Amount/Credit Limit'].fillna(0).astype(float)
        tran_sort_df['Tran.Amount/Credit Limit'] = ('RM' + tran_sort_df['Tran.Amount/Credit Limit'].astype(str))
    tran_sort_df['Tran.Type'] = tran_sort_df['Tran.Type'].fillna(0).astype(int)
    tran_sort_df['Tran.Type'] = (tran_sort_df['Tran.Type'].astype(str))
    tran_sort_df['Tran.Type'] = (tran_sort_df['Tran.Type'].str.rjust(3, "0"))
    tran_sort_df['Tran.Type']  = tran_sort_df['Tran.Type'].replace("000",'Null')
    tran_sort_df['Date (dd/mm/yyy)'] = pd.to_datetime(tran_sort_df['Date (dd/mm/yyy)'])
    tran_sort_df.fillna("Null", inplace=True)
    if (bank_list[0] == "1"):
        tran_sort_df = tran_sort_df[tran_sort_df['Credit Card Number'].str.startswith("001", na = False)]
    elif(bank_list[0] == "2"):
        tran_sort_df = tran_sort_df[tran_sort_df['Credit Card Number'].str.startswith("002", na = False)]
    elif (bank_list [0] == "3"):
        tran_sort_df = tran_sort_df[tran_sort_df['Credit Card Number'].str.startswith("003", na = False)]
    sort1=tran_sort_df[tran_sort_df["Status"]== "Completed" ]
    sort_1=sort1[sort1['Tran.Type']==(choose_list[0])]
    sort_2=sort_1[sort_1['Date (dd/mm/yyy)'].dt.year == (choose_list[1])]
    sort_3=sort_2[sort_2['Date (dd/mm/yyy)'].dt.month == (choose_list[2])]
    b1 = (len(sort_3)/20)
    b2 = (len(sort_3)//20)
    if b1>b2:
        sort_pages=b2+1
    elif b1 == b2 :
        sort_pages=b2
    elif b2 == 0 or b1 == 0 :
        sort_pages = 1
    if (len(sort_3)==0):
        print("\n")
        print("                                           ABC Company                                         ", datetime.date.today())
        print("-"*120)
        print("Report   : Transactions Sorted Report ")
        print("Transaction Type :",choose_list[0]," = ",transaction_type.get(choose_list[0][2]))
        print("Address  :  122,Jalan Ikan  , 20111 Kuala Kangsa , Malaysia")
        print("Year     :",choose_list[1])
        print("Month    :",choose_list[2],"   "*21,"[ Page = 1 ] ")
        print("-"*120)
        print("                    [Empty Record]            ")
        print(" "*80,"(Q)uit")
        print("-"*120)
        option=input("Option = ").upper()
        while (option != "Q"):
            print ("\t[Please enter again.]")
            option=input("Option = ").upper()
        if option == "Q" :
            choose_list.clear()
            tran_sort_option( )
    else:
        n = (len(sort_3)//sort_pages)
        group_n.append(0)
        if (isinstance(n,int)) == True :
            for i in range (sort_pages):
                group_n.append(n)
        else:
            for i in range (sort_pages):
                if (sort_pages - i) != 1:
                    n= (len(sort_3)//sort_pages)
                    group_n.append(n)
                else:
                    special_n=(len(sort_3)-(i*n))
                    group_n.append(special_n)
        groups = np.cumsum(group_n)
        d = {}
        for i in range(sort_pages):
            print("\n")
            print("                                           ABC Company                                         ", datetime.date.today())
            print("-"*120)
            print("Report   : Transactions Sorted Report ")
            print("Transaction Type :",choose_list[0]," = ",transaction_type.get(choose_list[0][2]))
            print("Address  :  122,Jalan Ikan  , 20111 Kuala Kangsa , Malaysia")
            print("Year     :",choose_list[1])
            print("Month    :",choose_list[2],"   "*21,"[ Page",sort_pages,"of",i+1,"]")
            print("-"*120)
            print("\n")
            d[i+1] = sort_3.iloc[groups[i]:groups[i+1]]
            print(d[i+1])
            count=0
            count_customer =0
            print("\n")
            if (choose_list[0] == "001") or (choose_list[0]== "002") or (choose_list[0]=="003")or (choose_list[0] == "004")or (choose_list[0]== "008"):
                for i in range (len(tran_type_list)):
                    if ((int(choose_list[0][2]) == tran_type_list[i]) and (tran_status_list[i]== "Completed") and (tran_year_list[i] == choose_list[1]) and (tran_month_list[i] == choose_list[2]) and (tran_credit_card_list[i][2] == str(bank_list[0]))):
                        count+=(tran_amount_list[i])
                        count_customer+=1
                print("\n")
                print(" "*90 ,"Total amount = RM ",count)
                print(" "*90,"Total of Customer  = ",count_customer)
                count=0
                count_customer =0
            if (choose_list[0] == "007") or (choose_list[0] == "009") or (choose_list[0] == "006"):
                for i in range (len(tran_type_list)):
                    if ((int(choose_list[0][2]) == tran_type_list[i]) and (tran_status_list[i]== "Completed") and (tran_year_list[i] == choose_list[1]) and (tran_month_list[i] == choose_list[2])and (tran_credit_card_list[i][2] == str(bank_list[0]))):
                        count_customer +=1
                print("\n")
                print(" "*90,"Total of Customer = ",count_customer)
                count=0
                count_customer=0
            print("-"*120)
            print(" "*80,"(N)ext   (P)rint   (Q)uit")
            print("-"*120)
            option=input("Option = ").upper()
            while (option!="N") and (option != "Q") and (option != "P"):
                print ("\t[Please enter again.]")
                option=input("Option = ")
            if (option== "N"):
                if ((sort_pages - i) != 1)and (sort_pages != 1) :
                    os.system("cls")
                    continue
                if (sort_pages == 1 ):
                    tran_sort_option_display(tran_file,column)
                else:
                    tran_sort_option_display(tran_file,column)
            if (option== "P"):
                with open ("display.txt" , "r") as f :
                    for line in f.readlines()[25:33]:
                        print(line)
                option=input("Option = ")
                option_list=["1","2","3","4"]
                while (option not in option_list):
                    print ("\t[Please enter again.]")
                    option=input("Option = ")
                sort_filename="Transaction Sorted File_"+""+str(time.strftime("%Y%m%d-%H%M%S"))
                if (option =="1"):
                    print("...................Processing...................")
                    print(" .................Please Wait..................")
                    sort_filename= sort_filename + ".txt"
                    with open (sort_filename,'w') as f :
                        for i in range(sort_pages):
                            f.write("\n\n")
                            f.write("                                           ABC Company                                         ")
                            f.write(str(datetime.date.today()))
                            f.write(str(time.strftime("%H:%M:%S")))
                            f.write("\n")
                            f.write("-"*120)
                            f.write("\n")
                            f.write("Report           : Transactions Sorted Report \n")
                            f.write("Transaction Type :")
                            f.write(choose_list[0])
                            f.write(" = ")
                            f.write(str(transaction_type.get(choose_list[0][2])))
                            f.write("\n")
                            f.write("Address  :  122,Jalan Ikan , 20111 Kuala Kangsa , Malaysia\n")
                            f.write("Year     : ")
                            f.write(str(choose_list[1]))
                            f.write("\n")
                            f.write("Month    : ")
                            f.write(str(choose_list[2]))
                            f.write("   "*21)
                            f.write("Page ")
                            f.write(str(sort_pages))
                            f.write(" of ")
                            f.write(str(i+1))
                            f.write("\n")
                            f.write("-"*120)
                            f.write("\n")
                            f.write(str(d[i+1]))
                            f.write("\n")
                            if (choose_list[0] == "001") or (choose_list[0]== "002") or (choose_list[0]=="003")or (choose_list[0] == "004")or (choose_list[0]== "008"):
                                for i in range (len(tran_type_list)):
                                    if ((int(choose_list[0][2]) == tran_type_list[i]) and (tran_status_list[i]== "Completed") and (tran_year_list[i] == int(choose_list[1])) and (int(tran_month_list[i]) == choose_list[2])and (int(tran_credit_card_list[i][2]) == int(bank_list[0]))):
                                        count+=(tran_amount_list[i])
                                        count_customer+=1
                                f.write("\n")
                                f.write(" "*90)
                                f.write("Total amount = RM ")
                                f.write(str(count))
                                f.write("\n")
                                f.write(" "*90)
                                f.write("Total of Customer = ")
                                f.write(str(count_customer))
                                f.write("\n")
                                f.write("-"*120)
                            count=0
                            count_customer =0
                            if (choose_list[0] == "007") or (choose_list[0] == "009") or (choose_list[0] == "006"):
                                for i in range (len(tran_type_list)):
                                    if ((int(choose_list[0][2]) == tran_type_list[i]) and (tran_status_list[i]== "Completed") and (tran_year_list[i] == choose_list[1]) and (tran_month_list[i] == choose_list[2])and (tran_credit_card_list[i][2] == str(bank_list[0]))):
                                        count_customer +=1
                                f.write("\n")
                                f.write(" "*90)
                                f.write("Total of Customer = ")
                                f.write(str(count_customer))
                                f.write("\n")
                                f.write("-"*120)
                            count=0
                            count_customer=0
                    print("Please Check your file >>> File Name = ",sort_filename)
                    print("Location : Documents // this File \n Thank You")
                    enter=input("Press Enter to continue :)")
                    group_n.clear()
                    tran_sort_option_display(tran_file,column)
                elif (option =="2")or (option == "3"):
                    print("...................Processing...................")
                    print(" .................Please Wait..................")
                    from docx import Document
                    from docx.shared import Inches
                    word_sort_filename= sort_filename + ".docx"
                    document = Document()
                    document.add_heading("                        [ABC Company]         ",0).bold = True
                    document.add_paragraph("Report   : Transactions Sorted Report ").bold=True
                    document.add_paragraph("Address  :  122,Jalan Ikan , 20111 Kuala Kangsa , Malaysia")
                    document.add_paragraph("Reported by : Alex Pan Wan")
                    document.add_paragraph("_________________________________________________________________________________________________________")
                    document.add_paragraph(str(datetime.date.today()))
                    document.add_paragraph(str(time.strftime("%H:%M:%S")))
                    t = document.add_table(sort_3.shape[0]+1, sort_3.shape[1])
                    for j in range(sort_3.shape[-1]):
                        t.cell(0,j).text = sort_3.columns[j]
                    for i in range (sort_3.shape[0]):
                        for j in range(sort_3.shape[-1]):
                            t.cell(i+1,j).text = str (sort_3.values[i,j])
                    document.save(str(word_sort_filename))
                    if (option == "2"):
                        print("Please Check your file >>> File Name = ",word_sort_filename)
                        enter=input("Press Enter to continue :)")
                        group_n.clear()
                        tran_sort_option_display(tran_file,column)
                    elif (option == "3"):
                        from docx2pdf import convert
                        pdf_sort_filename= sort_filename + ".pdf"
                        convert(word_sort_filename,pdf_sort_filename)
                        os.remove(word_sort_filename)
                        print("Please Check your file >>> File Name = ",pdf_sort_filename)
                        print("Location : Documents // this File \n Thank You")
                        enter=input("Press Enter to continue :)")
                        group_n.clear()
                        tran_sort_option_display(tran_file,column)
                elif(option == "4"):
                    group_n.clear()
                    tran_sort_option_display(tran_file,column)
            if (option== "Q"):
                choose_list.clear()
                tran_sort_option( )



### A.1.2.2 - Sort Option       = 1. Transaction Type + Month + Year
###         - Transaction Type  = 10
###         - Print Sorted Credit Card Transaction Record + Create text file // PDF file // WORD file (No Error)
## This function will display all the transaction record that user choose based on transaction type + month + year and user can print the sorted transaction record . 
def tran_sort_option_display_10(tran_file,column):
    os.system("cls")
    global choose_list
    group_n =[]
    tran_year_list,tran_month_list ,tran_type_list,tran_merchant_id_list,tran_account_id_list,tran_status_list,tran_amount_list ,tran_credit_card_list = read_updated_transaction_file(transaction_file)
    correct_pages ,error_pages = count_page()
    tran_sort_df = pd.concat(pd.read_excel(tran_file, sheet_name=None,header = 2 ,usecols=column), ignore_index=True)
    tran_sort_df['Merchant Id number'] = tran_sort_df['Merchant Id number'].fillna(0).astype(int)
    tran_sort_df['Merchant Id number'] = (tran_sort_df['Merchant Id number'].astype(str))
    tran_sort_df['Merchant Id number'] = (tran_sort_df['Merchant Id number'].str.rjust(3, "0"))
    tran_sort_df['Merchant Id number']  = tran_sort_df['Merchant Id number'].replace("000",'Null')            
    tran_sort_df['Tran.Type'] = tran_sort_df['Tran.Type'].fillna(0).astype(int)
    tran_sort_df['Tran.Type'] = (tran_sort_df['Tran.Type'].astype(str))
    tran_sort_df['Tran.Type'] = (tran_sort_df['Tran.Type'].str.rjust(3, "0"))
    tran_sort_df['Tran.Type']  = tran_sort_df['Tran.Type'].replace("000",'Null')
    tran_sort_df['Tran.Amount/Credit Limit']= tran_sort_df['Tran.Amount/Credit Limit'].fillna(0).astype(float)
    tran_sort_df['Tran.Amount/Credit Limit'] = ('RM' + tran_sort_df['Tran.Amount/Credit Limit'].astype(str))
    tran_sort_df['Date (dd/mm/yyy)'] = pd.to_datetime(tran_sort_df['Date (dd/mm/yyy)'])
    tran_sort_df.fillna("Null", inplace=True)
    if (bank_list[0] == "1"):
        tran_sort_df = tran_sort_df[tran_sort_df['Credit Card Number'].str.startswith("001", na = False)]
    elif(bank_list[0] == "2"):
        tran_sort_df = tran_sort_df[tran_sort_df['Credit Card Number'].str.startswith("002", na = False)]
    elif (bank_list [0] == "3"):
        tran_sort_df = tran_sort_df[tran_sort_df['Credit Card Number'].str.startswith("003", na = False)]
    sort_1=tran_sort_df[tran_sort_df["Status"]== "Completed" ]
    sort_2=sort_1[sort_1['Date (dd/mm/yyy)'].dt.year == (choose_list[1])]
    sort_3=sort_2[sort_2['Date (dd/mm/yyy)'].dt.month == (choose_list[2])]
    b1 = (len(sort_3)/20)
    b2 = (len(sort_3)//20)
    if b1>b2:
        sort_pages=b2+1
    elif b1 == b2 :
        sort_pages=b2
    elif b2 == 0 or b1 == 0 :
        sort_pages = 1
    if (len(sort_3)==0):
        print("\n")
        print("                                           ABC Company                                         ", datetime.date.today())
        print("-"*120)
        print("Report   : Transactions Sorted Report ")
        print("Address  :  122,Jalan Ikan  , 20111 Kuala Kangsa , Malaysia")
        print("Year     :",choose_list[1])
        print("Month    :",choose_list[2],"   "*21,"[ Page = 1 ] ")
        print("-"*120)
        print("                    [Empty Record]            ")
        print(" "*80,"(Q)uit")
        print("-"*120)
        option=input("Option = ").upper()
        while (option != "Q"):
            print ("\t[Please enter again.]")
            option=input("Option = ").upper()
        if option == "Q" :
            choose_list.clear()
            tran_sort_option( )
    else:
        n = (len(sort_3)//sort_pages)
        group_n.append(0)
        if (isinstance(n,int)) == True :
            for i in range (sort_pages):
                group_n.append(n)
        else:
            for i in range (sort_pages):
                if (sort_pages - i) != 1:
                    n= (len(sort_3)//sort_pages)
                    group_n.append(n)
                else:
                    special_n=(len(sort_3)-(i*n))
                    group_n.append(special_n)
        groups = np.cumsum(group_n)
        d = {}
        for i in range(sort_pages):
            print("\n")
            print("                                           ABC Company                                         ", datetime.date.today())
            print("-"*120)
            print("Report   : Transactions Sorted Report ")
            print("Address  :  122,Jalan Ikan  , 20111 Kuala Kangsa , Malaysia")
            print("Year     :",choose_list[1])
            print("Month    :",choose_list[2],"   "*21,"[ Page",sort_pages,"of",i+1,"]")
            print("-"*120)
            print("\n")
            d[i+1] = sort_3.iloc[groups[i]:groups[i+1]]
            print(d[i+1])
            print("\n")
            print("\n","  "*60,"Total record =",(len(sort_3)),"\n")
            print("-"*120)
            print(" "*80,"(N)ext   (P)rint   (Q)uit")
            print("-"*120)
            option=input("Option = ").upper()
            while (option!="N") and (option != "Q") and (option != "P"):
                print ("\t[Please enter again.]")
                option=input("Option = ")
            if (option== "N"):
                if ((sort_pages - i) != 1)and (sort_pages != 1) :
                    continue
                if (sort_pages == 1 ):
                    tran_sort_option_display(tran_file,column)
                else:
                    tran_sort_option_display(tran_file,column)
            if (option== "P"):
                with open ("display.txt" , "r") as f :
                    for line in f.readlines()[25:33]:
                        print(line)
                option=input("Option = ")
                option_list=["1","2","3","4"]
                while (option not in option_list):
                    option=input("Option = ")
                    print ("\t[Please enter again.]")
                sort_filename="Transaction Sorted File_"+""+str(time.strftime("%Y%m%d-%H%M%S"))
                if (option =="1"):
                    print("...................Processing...................")
                    print(" .................Please Wait..................")
                    sort_filename= sort_filename + ".txt"
                    with open (sort_filename,'w') as f :
                        for i in range(sort_pages):
                            f.write("\n\n")
                            f.write("                                           ABC Company                                         ")
                            f.write(str(datetime.date.today()))
                            f.write(str(time.strftime("%H:%M:%S")))
                            f.write("\n")
                            f.write("-"*120)
                            f.write("\n")
                            f.write("Report           : Transactions Sorted Report \n")
                            f.write("\n")
                            f.write("Address  :  122,Jalan Ikan , 20111 Kuala Kangsa , Malaysia\n")
                            f.write("Year     : ")
                            f.write(str(choose_list[1]))
                            f.write("\n")
                            f.write("Month    : ")
                            f.write(str(choose_list[2]))
                            f.write("   "*21)
                            f.write("Page ")
                            f.write(str(sort_pages))
                            f.write(" of ")
                            f.write(str(i+1))
                            f.write("\n")
                            f.write("-"*120)
                            f.write("\n")
                            f.write(str(d[i+1]))
                            f.write("\n")
                            f.write("  "*60)
                            f.write("Total record =")
                            f.write(str(len(sort_3)))
                            f.write("\n")
                    print("Please Check your file >>> File Name = ",sort_filename)
                    print("Location : Documents // this File \n Thank You")
                    enter=input("Press Enter to continue :)")
                    group_n.clear()
                    tran_sort_option_display_10(tran_file,column)
                elif (option =="2")or (option == "3"):
                    print("...................Processing...................")
                    print(" .................Please Wait..................")
                    from docx import Document
                    from docx.shared import Inches
                    word_sort_filename= sort_filename + ".docx"
                    document = Document()
                    document.add_heading("                        [ABC Company]         ",0).bold = True
                    document.add_paragraph("Report   : Transactions Sorted  Report ").bold=True
                    document.add_paragraph("Address  :  122,Jalan Ikan , 20111 Kuala Kangsa , Malaysia")
                    document.add_paragraph("Reported by : Alex Pan Wan")
                    document.add_paragraph("_________________________________________________________________________________________________________")
                    document.add_paragraph(str(datetime.date.today()))
                    document.add_paragraph(str(time.strftime("%H:%M:%S")))
                    t = document.add_table(sort_3.shape[0]+1, sort_3.shape[1])
                    for j in range(sort_3.shape[-1]):
                        t.cell(0,j).text = sort_3.columns[j]
                    for i in range (sort_3.shape[0]):
                        for j in range(sort_3.shape[-1]):
                            t.cell(i+1,j).text = str (sort_3.values[i,j])
                    document.save(str(word_sort_filename))
                    if (option == "2"):
                        print("Please Check your file >>> File Name = ",word_sort_filename)
                        enter=input("Press Enter to continue :)")
                        group_n.clear()
                        tran_sort_option_display_10(tran_file,column)
                    elif (option == "3"):
                        from docx2pdf import convert
                        pdf_sort_filename= sort_filename + ".pdf"
                        convert(word_sort_filename,pdf_sort_filename)
                        os.remove(word_sort_filename)
                        print("Please Check your file >>> File Name = ",pdf_sort_filename)
                        print("Location : Documents // this File \n Thank You")
                        enter=input("Press Enter to continue :)")
                        group_n.clear()
                        tran_sort_option_display_10(tran_file,column)
                elif(option == "4"):
                    group_n.clear()
                    tran_sort_option_display(tran_file,column)
            if (option== "Q"):
                choose_list.clear()
                tran_sort_option( )



### A.2 - Sort Option   = 2. Customer Account (eg:00X-XXXX) + Month + Year
#   This function allows user to insert the customer Account ID + Month + Year 
def tran_sort_option1( ):
    os.system("cls")
    correct_pages ,error_pages = count_page()
    tran_year_list,tran_month_list ,tran_type_list,tran_merchant_id_list,tran_account_id_list,tran_status_list,tran_amount_list ,tran_credit_card_list = read_updated_transaction_file(transaction_file)
    acc_id_list,acc_name_list,acc_ic_list,acc_credit_list,acc_balance_list,acc_status_list,acc_address_list = read_updated_account_file(account_file)
    group_n=[]
    global choose_list
    global global_list
    global loop
    print("-------------------------------------------------------------------")
    print("                       Insert Customer's Account Id                ")
    print("-------------------------------------------------------------------")
    print("              Account ID         Name    Balance     Status")
    for i in range (len(acc_id_list)):
        acc_id_list[i]=str(acc_id_list[i])
        if acc_id_list[i][2] ==  bank_list[0]:
            print("          ","*",acc_id_list[i] ,"     ", acc_name_list[i],"        RM",acc_balance_list[i] ,"   ",acc_status_list[i])
    print("-------------------------------------------------------------------")
    acc_option=input("Insert Customer's Account Id (eg:00X-XXXX):")
    while ((acc_option not in acc_id_list) and (acc_option[2] == bank_list[0])):
        print ("\t[Please enter again.]")
        acc_option=input("Insert Customer's Account Id (eg:00X-XXXX):")
    choose_list.append(acc_option)
    print("-------------------------------------------------------------------")
    print("                        Choose Year")
    print("-------------------------------------------------------------------")
    space1_list = tran_year_list.copy()
    space1_list.sort()
    for i in range (len(space1_list)):
        if (space1_list [i] not in global_list):
            global_list.append(space1_list[i])
            print("                 ","*",space1_list[i])
    print("-------------------------------------------------------------------")
    while loop :
        try:
            year=input("Year (eg:20XX) = ")
            if (int(year) in tran_year_list):
                loop = False
            else:
                loop = True
                print ("\t[Please enter again.]")
        except ValueError:
            loop = True
            print ("\t[Please enter again.]")
    choose_list.append(int(year))
    global_list.clear()
    loop = True 
    print("-------------------------------------------------------------------")
    print("                        Choose Month")
    print("-------------------------------------------------------------------")
    space2_list = tran_month_list.copy()
    space2_list.sort()
    for i in range (len(space2_list)):
        if (space2_list [i] not in global_list):
            global_list.append(space2_list[i])
            print("                 *",space2_list[i],"=",calendar.month_name[space2_list[i]])
    print("-------------------------------------------------------------------")
    while loop :
        try:
            month=input("Month (eg:1) = ")
            if (int(month) in tran_month_list):
                loop = False
            else:
                loop = True
                print ("\t[Please enter again.]")
        except ValueError:
            loop = True
            print ("\t[Please enter again.]")
    loop = True 
    choose_list.append(int(month))
    global_list.clear()
    cols=[0,2,3,4,5,6,11,12]
    tran_sort_option1_display(transaction_file,cols)


     
### A.2.1 - Display and Print the Credit Card Transaction Record in TEXT file // WORD file // PDF file (No Error)
###       - Based on the customer's Account ID , Month , Year that user choose.
def tran_sort_option1_display(tran_file,column):
    os.system("cls")
    global choose_list
    group_n =[]
    acc_id_list,acc_name_list,acc_ic_list,acc_credit_list,acc_balance_list,acc_status_list,acc_address_list = read_updated_account_file(account_file)
    tran_year_list,tran_month_list ,tran_type_list,tran_merchant_id_list,tran_account_id_list,tran_status_list,tran_amount_list ,tran_credit_card_list = read_updated_transaction_file(transaction_file)
    correct_pages ,error_pages = count_page()
    position = acc_id_list.index(choose_list[0])
    tran_sort_df = pd.concat(pd.read_excel(tran_file, sheet_name=None,header = 2 ,usecols=column), ignore_index=True)
    tran_sort_df = pd.concat(pd.read_excel(tran_file, sheet_name=None,header = 2 ,usecols=column), ignore_index=True)
    tran_sort_df['Merchant Id number'] = tran_sort_df['Merchant Id number'].fillna(0).astype(int)
    tran_sort_df['Merchant Id number'] = (tran_sort_df['Merchant Id number'].astype(str))
    tran_sort_df['Merchant Id number'] = (tran_sort_df['Merchant Id number'].str.rjust(3, "0"))
    tran_sort_df['Merchant Id number']  = tran_sort_df['Merchant Id number'].replace("000",'Null')            
    tran_sort_df['Tran.Type'] = tran_sort_df['Tran.Type'].fillna(0).astype(int)
    tran_sort_df['Tran.Type'] = (tran_sort_df['Tran.Type'].astype(str))
    tran_sort_df['Tran.Type'] = (tran_sort_df['Tran.Type'].str.rjust(3, "0"))
    tran_sort_df['Tran.Type']  = tran_sort_df['Tran.Type'].replace("000",'Null')
    tran_sort_df['Tran.Amount/Credit Limit']= tran_sort_df['Tran.Amount/Credit Limit'].fillna(0).astype(float)
    tran_sort_df['Tran.Amount/Credit Limit'] = ('RM' + tran_sort_df['Tran.Amount/Credit Limit'].astype(str))
    tran_sort_df['Date (dd/mm/yyy)'] = pd.to_datetime(tran_sort_df['Date (dd/mm/yyy)'])
    tran_sort_df['Date (dd/mm/yyy)'] = pd.to_datetime(tran_sort_df['Date (dd/mm/yyy)'])
    tran_sort_df.fillna("Null", inplace=True)
    sort1=tran_sort_df[tran_sort_df["Status"]== "Completed" ]
    sort_1=sort1[sort1['Credit Card Number']==(choose_list[0])]
    sort_2=sort_1[sort_1['Date (dd/mm/yyy)'].dt.year == (choose_list[1])]
    sort_3=sort_2[sort_2['Date (dd/mm/yyy)'].dt.month == (choose_list[2])]
    b1 = (len(sort_3)/20)
    b2 = (len(sort_3)//20)
    if b1>b2:
        sort_pages=b2+1
    elif b1 == b2 :
        sort_pages=b2
    elif b2 == 0 or b1 == 0 :
        sort_pages = 1
        
    if (len(sort_3)==0):
        print("\n")
        print("                                           ABC Company                                         ", datetime.date.today())
        print("-"*120)
        print("Report   : Transactions Sorted Report ")
        print("Customer Name   :",acc_name_list[position])
        print("Customer IC     :",acc_ic_list[position])
        print("Customer Add    :",acc_address_list[position])
        print("Customer Balance:",acc_balance_list[position])
        print("Address  :  122,Jalan Ikan  , 20111 Kuala Kangsa , Malaysia")
        print("Year     :",choose_list[1])
        print("Month    :",choose_list[2],"   "*21,"[ Page = 1 ] ")
        print("-"*120)
        print("                    [Empty Record]            ")
        print(" "*80,"(Q)uit")
        print("-"*120)
        option=input("Option = ").upper()
        while (option != "Q"):
            print ("\t[Please enter again.]")
            option=input("Option = ").upper()
        if option == "Q" :
            choose_list.clear()
            transaction_report(correct_pages,transaction_file)
    else:
        n = (len(sort_3)//sort_pages)
        group_n.append(0)
        if (isinstance(n,int)) == True :
            for i in range (sort_pages):
                group_n.append(n)
        else:
            for i in range (sort_pages):
                if (sort_pages - i) != 1:
                    n= (len(sort_3)//sort_pages)
                    group_n.append(n)
                else:
                    special_n=(len(sort_3)-(i*n))
                    group_n.append(special_n)
        groups = np.cumsum(group_n)
        d = {}
        for i in range(sort_pages):
            print("\n")
            print("                                           ABC Company                                         ", datetime.date.today())
            print("-"*120)
            print("Report          : Transactions Sorted Report ")
            print("Customer ID     :",choose_list[0])
            print("Customer Name   :",acc_name_list[position])
            print("Customer IC     :",acc_ic_list[position])
            print("Customer Add    :",acc_address_list[position])
            print("Customer Balance:",acc_balance_list[position])
            print("Address         :  122,Jalan Ikan  , 20111 Kuala Kangsa , Malaysia")
            print("Year            :",choose_list[1])
            print("Month           :",choose_list[2],"   "*21,"[ Page",sort_pages,"of",i+1,"]")
            print("-"*120)
            print("\n")
            d[i+1] = sort_3.iloc[groups[i]:groups[i+1]]
            print(d[i+1])
            count=0
            count_customer =0
            print("\n")
            print("  "*60,"Total record =",(len(sort_3)),"\n")
            print("-"*120)
            print(" "*80,"(N)ext   (P)rint   (Q)uit")
            print("-"*120)
            option=input("Option = ").upper()
            while (option!="N") and (option != "Q") and (option != "P"):
                print ("\t[Please enter again.]")
                option=input("Option = ")
            if (option== "N"):
                if ((sort_pages - i) != 1)and (sort_pages != 1) :
                    os.system("cls")
                    continue
                if (sort_pages == 1 ):
                    tran_sort_option1_display(tran_file,column)
                else:
                    tran_sort_option1_display(tran_file,column)
            if (option== "P"):
                with open ("display.txt" , "r") as f :
                    for line in f.readlines()[25:33]:
                        print(line)
                option=input("Option = ")
                option_list=["1","2","3","4"]
                while (option not in option_list):
                    option=input("Option = ")
                    print ("\t[Please enter again.]")
                sort_filename="Transaction Sorted File_"+""+str(time.strftime("%Y%m%d-%H%M%S"))
                if (option =="1"):
                    print("...................Processing...................")
                    print(" .................Please Wait..................")
                    sort_filename= sort_filename + ".txt"
                    with open (sort_filename,'w') as f :
                        for i in range(sort_pages):
                            f.write("\n\n")
                            f.write("                                           ABC Company                                         ")
                            f.write(str(datetime.date.today()))
                            f.write(str(time.strftime("%H:%M:%S")))
                            f.write("\n")
                            f.write("-"*120)
                            f.write("\n")
                            f.write("Report           : Transactions Sorted Report ")
                            f.write("\n")
                            f.write("Customer ID   : ")
                            f.write(str(choose_list[0]))
                            f.write("\n")
                            f.write("Customer Name  : ")
                            f.write(str(acc_name_list[position]))
                            f.write("\n")
                            f.write("Customer IC   : ")
                            f.write(str(acc_ic_list[position]))
                            f.write("\n")
                            f.write("Customer Add   : ")
                            f.write(str(acc_address_list[position]))
                            f.write("\n")
                            f.write("Customer Balance : ")
                            f.write(str(acc_balance_list[position]))
                            f.write("\n")
                            f.write("Address  :  122,Jalan Ikan , 20111 Kuala Kangsa , Malaysia\n")
                            f.write("Year     : ")
                            f.write(str(choose_list[1]))
                            f.write("\n")
                            f.write("Month    : ")
                            f.write(str(choose_list[2]))
                            f.write("   "*21)
                            f.write("Page ")
                            f.write(str(sort_pages))
                            f.write(" of ")
                            f.write(str(i+1))
                            f.write("\n")
                            f.write("-"*120)
                            f.write("\n")
                            f.write(str(d[i+1]))
                            f.write("\n")
                            f.write("  "*60)
                            f.write("Total record =")
                            f.write(str(len(sort_3)))
                            f.write("\n")
                            count=0
                            count_customer=0
                    print("Please Check your file >>> File Name = ",sort_filename)
                    print("Location : Documents // this File \n Thank You")
                    enter=input("Press Enter to continue :)")
                    group_n.clear()
                    tran_sort_option1_display(tran_file,column)
                elif (option =="2")or (option == "3"):
                    print("...................Processing...................")
                    print(" .................Please Wait..................")
                    from docx import Document
                    from docx.shared import Inches
                    word_sort_filename= sort_filename + ".docx"
                    document = Document()
                    document.add_heading("                        [ABC Company]         ",0).bold = True
                    document.add_paragraph("Report   : Transactions Sorted Report ").bold=True
                    document.add_paragraph("Address  :  122,Jalan Ikan , 20111 Kuala Kangsa , Malaysia")
                    document.add_paragraph("Reported by : Alex Pan Wan")
                    document.add_paragraph("_________________________________________________________________________________________________________")
                    document.add_paragraph(str(datetime.date.today()))
                    document.add_paragraph(str(time.strftime("%H:%M:%S")))
                    t = document.add_table(sort_3.shape[0]+1, sort_3.shape[1])
                    for j in range(sort_3.shape[-1]):
                        t.cell(0,j).text = sort_3.columns[j]
                    for i in range (sort_3.shape[0]):
                        for j in range(sort_3.shape[-1]):
                            t.cell(i+1,j).text = str (sort_3.values[i,j])
                    document.save(str(word_sort_filename))
                    if (option == "2"):
                        print("Please Check your file >>> File Name = ",word_sort_filename)
                        enter=input("Press Enter to continue :)")
                        group_n.clear()
                        tran_sort_option1_display(tran_file,column)
                    elif (option == "3"):
                        from docx2pdf import convert
                        pdf_sort_filename= sort_filename + ".pdf"
                        convert(word_sort_filename,pdf_sort_filename)
                        os.remove(word_sort_filename)
                        print("Please Check your file >>> File Name = ",pdf_sort_filename)
                        print("Location : Documents // this File \n Thank You")
                        enter=input("Press Enter to continue :)")
                        group_n.clear()
                        tran_sort_option1_display(tran_file,column)
                elif(option == "4"):
                    group_n.clear()
                    tran_sort_option1_display(tran_file,column)
            if (option== "Q"):
                choose_list.clear()
                transaction_report(correct_pages,transaction_file)


##B.1 [ Transaction Error Report ] - Error Record
# This function allows user to view all the transaction error record  .
def transaction_error_report(e_pages,tran_file):
    os.system("cls")
    global calendar
    group_n=[]
    tran_year_list,tran_month_list ,tran_type_list,tran_merchant_id_list,tran_account_id_list,tran_status_list,tran_amount_list ,tran_credit_card_list = read_updated_transaction_file(transaction_file)
    correct_pages ,error_pages = count_page()
    cols=[0,2,3,4,5,6,11,12]
    tran_error_df = pd.concat(pd.read_excel(tran_file, sheet_name=None,header = 2 ,usecols=cols), ignore_index=True)
    tran_error_df['Merchant Id number'] = tran_error_df['Merchant Id number'].fillna(0).astype(int)
    tran_error_df['Merchant Id number'] = (tran_error_df['Merchant Id number'].astype(str))
    tran_error_df['Merchant Id number'] = (tran_error_df['Merchant Id number'].str.rjust(3, "0"))
    tran_error_df['Merchant Id number']  = tran_error_df['Merchant Id number'].replace("000",'Null')            
    tran_error_df['Tran.Type'] = tran_error_df['Tran.Type'].fillna(0).astype(int)
    tran_error_df['Tran.Type'] = (tran_error_df['Tran.Type'].astype(str))
    tran_error_df['Tran.Type'] = (tran_error_df['Tran.Type'].str.rjust(3, "0"))
    tran_error_df['Tran.Type']  = tran_error_df['Tran.Type'].replace("000",'Null')
    tran_error_df['Tran.Amount/Credit Limit']= tran_error_df['Tran.Amount/Credit Limit'].fillna(0).astype(float)
    tran_error_df['Tran.Amount/Credit Limit'] = ('RM' + tran_error_df['Tran.Amount/Credit Limit'].astype(str))
    tran_error_df['Date (dd/mm/yyy)'] = pd.to_datetime(tran_error_df['Date (dd/mm/yyy)'])
    tran_error_df.fillna("Null", inplace=True)
    if (bank_list[0] == "1"):
        tran_error_df = tran_error_df[tran_error_df['Credit Card Number'].str.startswith("001", na = False)]
    elif(bank_list[0] == "2"):
        tran_error_df = tran_error_df[tran_error_df['Credit Card Number'].str.startswith("002", na = False)]
    elif (bank_list [0] == "3"):
        tran_error_df = tran_error_df[tran_error_df['Credit Card Number'].str.startswith("003", na = False)]
    pd.set_option('display.max_rows', None)
    pd.set_option('display.max_columns', None)
    pd.set_option('display.width', None)
    pd.set_option('display.max_colwidth', None)
    tran_error_df2=tran_error_df[tran_error_df["Status"]== "Failed" ]
    if (len(tran_error_df2) == 0):
        print("\n")
        print("                                           ABC Company                                 ", datetime.date.today() ,time.strftime("%H:%M:%S"))
        print("-"*120)
        print("Report   : Transactions Error Report (ALL) ")
        print("Address  :  122,Jalan Ikan  , 20111 Kuala Kangsa , Malaysia")
        print("Year     :",(min(tran_year_list)),"to",(max(tran_year_list)))
        print("Month    :",calendar.month_name[(min(tran_month_list))],"to",calendar.month_name[(max(tran_month_list))],"   "*21,"[ Page = 1 ]")
        print("-"*120)
        print("\n\n")
        print("                    [Empty Record]            ")
        print(" "*80,"(Q)uit")
        print("-"*120)
        option=input("Option = ").upper() 
        while (option != "Q"):
            print ("\t[Please enter again.]")
            option=input("Option = ").upper() 
        if option == "Q" :
            choose_list.clear()
            report_option()       
    else:
        n= (len(tran_error_df2)//e_pages)
        group_n.append(0)
        if (isinstance(n,int)) == True :
            for i in range (e_pages):
                group_n.append(n)
        else:
            for i in range (e_pages):
                if (e_pages - i) != 1:
                    n= (len(tran_df2)//e_pages)
                    group_n.append(n)
                else:
                    special_n=(len(tran_error_df2)-(i*n))
                    group_n.append(special_n)
        groups = np.cumsum(group_n)
        d = {}
        for i in range(e_pages):
            print("\n")
            print("                                           ABC Company                                 ", datetime.date.today() ,time.strftime("%H:%M:%S"))
            print("-"*120)
            print("Report   : Transactions Error Report  (ALL) ")
            print("Address  :  122,Jalan Ikan  , 20111 Kuala Kangsa , Malaysia")
            print("Year     :",(min(tran_year_list)),"to",(max(tran_year_list)))
            print("Month    :",calendar.month_name[(min(tran_month_list))],"to",calendar.month_name[(max(tran_month_list))],"   "*21,"[ Page",e_pages,"of",i+1,"]")
            print("-"*120)
            print("\n\n")
            d[i+1] = tran_error_df2.iloc[groups[i]:groups[i+1]]
            print(d[i+1])
            print("\n","  "*60,"Total record =",(len(tran_error_df2)),"\n")
            print(" "*70,"           (N)ext   (S)ort   (P)rint   (Q)uit")
            print("-"*120)
            option=input("Option = ").upper()
            option_list=["N","S","P","Q"]
            while option not in option_list:
                print ("\t[Please enter again.]")
                option=input("Option = ").upper()
            if(option == "N"):
                if (e_pages - i) != 1:
                    os.system("cls")
                    continue
                else:
                    transaction_error_report(e_pages,tran_file)
            elif (option == "S"): 
                group_n.clear()
                sort_option_list=["1","2","3"]
                with open ("display.txt" , "r") as f :
                    for line in f.readlines()[45:52]:
                        print(line)
                sort_option=input("Option = ")
                while (sort_option not in sort_option_list) :
                    print ("\t[Please enter again.]")
                    sort_option=input("Option = ")
                if (sort_option == "1"):
                    tran_error_sort_option( )
                elif (sort_option == "2"):
                    transaction_error_sort_option1()
                elif (sort_option == "3"):
                    transaction_error_report(e_pages,tran_file)
            elif (option == "P"):
                group_n.clear()
                tran_error_print_option(tran_error_df,tran_error_df2,error_pages)
            elif (option == "Q"):
                group_n.clear()
                report_option()
    return tran_error_df,tran_error_df2


### B.1.1 - Print the Credit Card Transaction Record in TEXT file // WORD file // PDF file (Error)
def tran_error_print_option(data1,data,c_pages):
    os.system("cls")
    group_n=[]
    tran_year_list,tran_month_list ,tran_type_list,tran_merchant_id_list,tran_account_id_list,tran_status_list,tran_amount_list ,tran_credit_card_list = read_updated_transaction_file(transaction_file)
    correct_pages ,error_pages = count_page()
    with open ("display.txt" , "r") as f :
        for line in f.readlines()[25:33]:
            print(line)
    option=input("Option = ")
    filename="Transaction Error File_"+""+str(time.strftime("%Y%m%d-%H%M%S"))
    option_list=["1","2","3","4"]
    while option not in option_list:
        print ("\t[Please enter again.]")
        option=input("Option = ")
    if (option =="1"):
        filename=filename + ".txt"        
        n= (len(data)//c_pages)
        group_n.append(0)
        if (isinstance(n,int)) == True :
            for i in range (c_pages):
                group_n.append(n)
        else:
            for i in range (c_pages):
                if (c_pages - i) != 1:
                    n= (len(data)//c_pages)
                    group_n.append(n)
                else:
                    special_n=(len(data)-(i*n))
                    group_n.append(special_n)
        groups = np.cumsum(group_n)
        with open (filename,"w") as f :
            for i in range(c_pages):
                d = {}
                d[i+1]= data.iloc[groups[i]:groups[i+1]]
                f.write("\n\n\n")
                f.write("                                           ABC Company                                         ")
                f.write(str(datetime.date.today()))
                f.write(str(time.strftime("%H:%M:%S")))
                f.write("\n")
                f.write("-"*120)
                f.write("\n")
                f.write("Report   : Transactions Error Report (ALL) \n")
                f.write("Transaction Type : All")
                f.write("Address  :  122,Jalan Ikan , 20111 Kuala Kangsa , Malaysia\n")
                f.write("Reported by  : Alex Pan Wan ")
                f.write("Year     : ")
                f.write(str((min(tran_year_list))))
                f.write(" to ")
                f.write(str((max(tran_year_list))))
                f.write("\n")
                f.write("Month    : ")
                f.write(str(calendar.month_name[(min(tran_month_list))]))
                f.write(" to ")
                f.write(str(calendar.month_name[(max(tran_month_list))]))
                f.write("   "*21)
                f.write("Page ")
                f.write(str(c_pages))
                f.write(" of ")
                f.write(str(i+1))
                f.write("\n")
                f.write("-"*120)
                f.write("\n")
                f.write(str(d[i+1]))
                f.write("\n\n")
                f.write("  "*60)
                f.write("Total record = ")
                f.write(str(len(data)))
                f.write("\n")
                f.write("-"*120)
                f.write("\n\n\n")
        print("Please Check your file >>> File Name = ",filename)
        print("Location : Documents // this File \n\n .......Thank You.....")
        enter=input("Press Enter to continue :)")
        group_n.clear()
        tran_error_print_option(data1,data,c_pages)
    elif (option =="2") or (option == "3"):
        print("...................Processing...................")
        print(" .................Please Wait..................")
        from docx import Document
        from docx.shared import Inches
        word_filename= filename + ".docx"
        document = Document()
        document.add_heading("                        [ABC Company]         ",0).bold = True
        document.add_paragraph("Report   : Transactions Error Report (ALL) ").bold=True
        document.add_paragraph("Address  :  122,Jalan Ikan , 20111 Kuala Kangsa , Malaysia")
        document.add_paragraph("Reported by : Alex Pan Wan")
        document.add_paragraph("_________________________________________________________________________________________________________")
        document.add_paragraph(str(datetime.date.today()))
        document.add_paragraph(str(time.strftime("%H:%M:%S")))
        t = document.add_table(data.shape[0]+1, data.shape[1])
        for j in range(data.shape[-1]):
            t.cell(0,j).text = data.columns[j]
        for i in range (data.shape[0]):
            for j in range(data.shape[-1]):
                t.cell(i+1,j).text = str (data.values[i,j])
        document.save(str(word_filename))
        if (option == "2"):
            print("Please Check your file >>> File Name = ",word_filename)
            enter=input("Press Enter to continue :)")
            group_n.clear()
            tran_error_print_option(data1,data,c_pages)
        elif (option == "3"):
            from docx2pdf import convert
            pdf_filename=filename + ".pdf"
            convert(word_filename,pdf_filename)
            os.remove(word_filename)
            print("Please Check your file >>> File Name = ",pdf_filename)
            print("Location : Documents // this File \n\n............ Thank You.............")
            enter=input("Press Enter to continue :)")
            group_n.clear()
            tran_error_print_option(data1,data,c_pages)
    elif (option == "4"):
        group_n.clear()
        transaction_error_report(error_pages,transaction_file)


### B.1.2 - Sort Option = 1. Transaction Type + Month + Year
##This function allows user to select transaction type + month + year .
def tran_error_sort_option( ):
    os.system("cls")
    option_list=["1","2","3","4","6","7","8","9","10","11"]
    tran_year_list,tran_month_list ,tran_type_list,tran_merchant_id_list,tran_account_id_list,tran_status_list,tran_amount_list ,tran_credit_card_list = read_updated_transaction_file(transaction_file)
    correct_pages ,error_pages = count_page()
    group_n=[]
    global choose_list
    global global_list
    global loop
    with open ("display.txt" , "r") as f :
        for line in f.readlines()[32:46]:
            print(line)
    option=input("Option = ")
    while (option not in option_list):
        print ("\t[Please enter again.]")
        option=input("Option = ")
    if (option != "10")and (option !="11"):
        t_type = "00" + option
        choose_list.append(t_type)
    if (option == "10"):
        t_type= "0" + option
        choose_list.append(t_type)
    if (option != "11"):
        print("-------------------------------------------------------------------")
        print("                        Choose Year")
        print("-------------------------------------------------------------------")
        space1_list = tran_year_list.copy()
        space1_list.sort()
        for i in range (len(space1_list)):
            if (space1_list [i] not in global_list):
                global_list.append(space1_list[i])
                print("                 ","*",space1_list[i])
        print("-------------------------------------------------------------------")
        while loop :
            try:
                year=input("Year (eg:20XX) = ")
                if (int(year) in tran_year_list):
                    loop = False
                else:
                    loop = True
                    print ("\t[Please enter agin.]")
            except ValueError:
                loop = True
                print ("\t[Please enter again.]")
        choose_list.append(int(year))
        global_list.clear()
        loop = True 
        print("-------------------------------------------------------------------")
        print("                        Choose Month")
        print("-------------------------------------------------------------------")
        space2_list = tran_month_list.copy()
        space2_list.sort()
        for i in range (len(space2_list)):
            if (space2_list [i] not in global_list):
                global_list.append(space2_list[i])
                print("                 *",space2_list[i],"=",calendar.month_name[space2_list[i]])
        print("-------------------------------------------------------------------")
        while loop :
            try:
                month=input("Month (eg:1) = ")
                if (int(month) in tran_month_list):
                    loop = False
                else:
                    loop = True
                    print ("\t[Please enter again.]")
            except ValueError:
                loop = True
                print ("\t[Please enter again.]")
        loop = True 
        choose_list.append(int(month))
        global_list.clear()
        if (choose_list[0] == "001") or (choose_list[0]=="004"):
            cols=[0,2,3,4,5,11,12]
            tran_error_sort_option_display(transaction_file,cols)
        elif (choose_list[0] == "002" ):
            cols=[0,2,3,4,11,12]
            tran_error_sort_option_display(transaction_file,cols)
        elif (choose_list[0] == "003"):
            cols=[0,2,3,4,6,11,12]
            tran_error_sort_option_display(transaction_file,cols)
        elif (choose_list[0] == "007" ) or (choose_list[0] == "009") or (choose_list[0]=="006"):
            cols=[0,2,3,11,12]
            tran_error_sort_option_display(transaction_file,cols)
        elif (choose_list[0] == "008"):
            cols=[0,2,3,4,5,11,12]
            tran_error_sort_option_display(transaction_file,cols)
        elif (choose_list[0] == "010"):
            cols=[0,2,3,4,5,6,7,8,11,12]
            tran_error_sort_option_display_10(transaction_file,cols)
    else:
        group_n.clear()
        choose_list.clear()
        transaction_error_report(error_pages,transaction_file)



### B.1.2.1 - Print Sorted Credit Card Transaction Record + Create text file // PDF file // WORD file (Error)
###       - Transac tion Type = 1,2,3,4,6,7,8,9
## This function will display all the transaction error record that user choose based on transaction type + month + year and user can print the sorted transaction record .
def tran_error_sort_option_display(tran_file,column):
    os.system("cls")
    global choose_list
    group_n =[]
    tran_year_list,tran_month_list ,tran_type_list,tran_merchant_id_list,tran_account_id_list,tran_status_list,tran_amount_list ,tran_credit_card_list = read_updated_transaction_file(transaction_file)
    correct_pages ,error_pages = count_page()
    transaction_type={"1": "Purchase (Physical)" , "2" : "Cash Withdrawal", "3": "Payment/Deposit into Account","4":"Purchase (Online)"  , "6" : "change of address", "7" : "change PIN numbers" , "8" : "Open/Approval of New Account" , "9": "Close an Account" }
    tran_sort_df = pd.concat(pd.read_excel(tran_file, sheet_name=None,header = 2 ,usecols=column), ignore_index=True)
    if (choose_list[0][2] == "8") or (choose_list[0][2] == "1") or (choose_list[0][2]=="4"):
        tran_sort_df['Merchant Id number'] = tran_sort_df['Merchant Id number'].fillna(0).astype(int)
        tran_sort_df['Merchant Id number'] = (tran_sort_df['Merchant Id number'].astype(str))
        tran_sort_df['Merchant Id number'] = (tran_sort_df['Merchant Id number'].str.rjust(3, "0"))
        tran_sort_df['Merchant Id number']  = tran_sort_df['Merchant Id number'].replace("000",'Null')
        tran_sort_df['Tran.Amount/Credit Limit']= tran_sort_df['Tran.Amount/Credit Limit'].fillna(0).astype(float)
        tran_sort_df['Tran.Amount/Credit Limit'] = ('RM' + tran_sort_df['Tran.Amount/Credit Limit'].astype(str))
    tran_sort_df['Tran.Type'] = tran_sort_df['Tran.Type'].fillna(0).astype(int)
    tran_sort_df['Tran.Type'] = (tran_sort_df['Tran.Type'].astype(str))
    tran_sort_df['Tran.Type'] = (tran_sort_df['Tran.Type'].str.rjust(3, "0"))
    tran_sort_df['Tran.Type']  = tran_sort_df['Tran.Type'].replace("000",'Null')
    tran_sort_df['Date (dd/mm/yyy)'] = pd.to_datetime(tran_sort_df['Date (dd/mm/yyy)'])
    tran_sort_df.fillna("Null", inplace=True)
    sort1=tran_sort_df[tran_sort_df["Status"]== "Failed" ]
    if (bank_list[0] == "1"):
        sort1 =sort1[sort1['Credit Card Number'].str.startswith("001", na = False)]
    elif(bank_list[0] == "2"):
        sort1 =sort1[sort1['Credit Card Number'].str.startswith("002", na = False)]
    elif (bank_list [0] == "3"):
        sort1 =sort1[sort1['Credit Card Number'].str.startswith("003", na = False)]
    sort_1=sort1[sort1['Tran.Type']==(choose_list[0])]
    sort_2=sort_1[sort_1['Date (dd/mm/yyy)'].dt.year == (choose_list[1])]
    sort_3=sort_2[sort_2['Date (dd/mm/yyy)'].dt.month == (choose_list[2])]
    b1 = (len(sort_3)/20)
    b2 = (len(sort_3)//20)
    if b1>b2:
        sort_pages=b2+1
    elif b1 == b2 :
        sort_pages=b2
    elif b2 == 0 or b1 == 0 :
        sort_pages = 1
    if (len(sort_3)==0):
        print("\n")
        print("                                           ABC Company                                         ", datetime.date.today())
        print("-"*120)
        print("Report   : Transactions Error Sorted Report ")
        print("Transaction Type :",choose_list[0]," = ",transaction_type.get(choose_list[0][2]))
        print("Address  :  122,Jalan Ikan  , 20111 Kuala Kangsa , Malaysia")
        print("Year     :",choose_list[1])
        print("Month    :",choose_list[2],"   "*21,"[ Page = 1 ] ")
        print("-"*120)
        print("                    [Empty Record]            ")
        print(" "*80,"(Q)uit")
        print("-"*120)
        option=input("Option = ").upper()
        while (option != "Q"):
            print ("\t[Please enter again.]")
            option=input("Option = ").upper()
        if option == "Q" :
            choose_list.clear()
            tran_error_sort_option( )

    else:
        n = (len(sort_3)//sort_pages)
        group_n.append(0)
        if (isinstance(n,int)) == True :
            for i in range (sort_pages):
                group_n.append(n)
        else:
            for i in range (sort_pages):
                if (sort_pages - i) != 1:
                    n= (len(sort_3)//sort_pages)
                    group_n.append(n)
                else:
                    special_n=(len(sort_3)-(i*n))
                    group_n.append(special_n)
        groups = np.cumsum(group_n)
        d = {}
        for i in range(sort_pages):
            print("\n")
            print("                                           ABC Company                                         ", datetime.date.today())
            print("-"*120)
            print("Report   : Transactions  Error Sorted Report ")
            print("Transaction Type :",choose_list[0]," = ",transaction_type.get(choose_list[0][2]))
            print("Address  :  122,Jalan Ikan  , 20111 Kuala Kangsa , Malaysia")
            print("Year     :",choose_list[1])
            print("Month    :",choose_list[2],"   "*21,"[ Page",sort_pages,"of",i+1,"]")
            print("-"*120)
            print("\n")
            d[i+1] = sort_3.iloc[groups[i]:groups[i+1]]
            print(d[i+1])
            count=0
            count_customer =0
            print("\n")
            if (choose_list[0] == "001") or (choose_list[0]== "002") or (choose_list[0]=="003")or (choose_list[0] == "004")or (choose_list[0]== "008"):
                for i in range (len(tran_type_list)):
                    if ((int(choose_list[0][2]) == tran_type_list[i]) and (tran_status_list[i]== "Failed") and (tran_year_list[i] == choose_list[1]) and (tran_month_list[i] == choose_list[2])and (tran_credit_card_list[i][2] == str(bank_list[0]))):
                        count+=(tran_amount_list[i])
                        count_customer+=1
                print("\n")
                print(" "*90 ,"Total amount = RM ",count)
                print(" "*90,"Total of Customer  = ",count_customer)
                count=0
                count_customer =0
            if (choose_list[0] == "007") or (choose_list[0] == "009") or (choose_list[0] == "006"):
                for i in range (len(tran_type_list)):
                    if ((int(choose_list[0][2]) == tran_type_list[i]) and (tran_status_list[i]== "Failed") and (tran_year_list[i] == choose_list[1]) and (tran_month_list[i] == choose_list[2])and (tran_credit_card_list[i][2] == str(bank_list[0]))):
                        count_customer +=1
                print("\n")
                print(" "*90,"Total of Customer = ",count_customer)
                count=0
                count_customer=0
            print("-"*120)
            print(" "*80,"(N)ext   (P)rint   (Q)uit")
            print("-"*120)
            option=input("Option = ").upper()
            while (option!="N") and (option != "Q") and (option != "P"):
                print ("\t[Please enter again.]")
                option=input("Option = ")
            if (option== "N"):
                if ((sort_pages - i) != 1)and (sort_pages != 1) :
                    os.system("cls")
                    continue
                if (sort_pages == 1 ):
                    tran_error_sort_option_display(tran_file,column)
                else:
                    tran_error_sort_option_display(tran_file,column)
            if (option== "P"):
                with open ("display.txt" , "r") as f :
                    for line in f.readlines()[25:33]:
                        print(line)
                option=input("Option = ")
                option_list=["1","2","3","4"]
                while (option not in option_list):
                    option=input("Option = ")
                    print ("\t[Please enter again.]")
                sort_filename="Transaction Error Sorted File_"+""+str(time.strftime("%Y%m%d-%H%M%S"))
                if (option =="1"):
                    print("...................Processing...................")
                    print(" .................Please Wait..................")
                    sort_filename= sort_filename + ".txt"
                    with open (sort_filename,'w') as f :
                        for i in range (sort_pages):
                            d[i+1] = sort_3.iloc[groups[i]:groups[i+1]]
                            f.write("\n\n")
                            f.write("                                           ABC Company                                         ")
                            f.write(str(datetime.date.today()))
                            f.write(str(time.strftime("%H:%M:%S")))
                            f.write("\n")
                            f.write("-"*120)
                            f.write("\n")
                            f.write("Report           : Transactions  Error Sorted Report \n")
                            f.write("Transaction Type :")
                            f.write(choose_list[0])
                            f.write(" = ")
                            f.write(str(transaction_type.get(choose_list[0][2])))
                            f.write("\n")
                            f.write("Address  :  122,Jalan Ikan , 20111 Kuala Kangsa , Malaysia\n")
                            f.write("Year     : ")
                            f.write(str(choose_list[1]))
                            f.write("\n")
                            f.write("Month    : ")
                            f.write(str(choose_list[2]))
                            f.write("   "*21)
                            f.write("Page ")
                            f.write(str(sort_pages))
                            f.write(" of ")
                            f.write(str(i+1))
                            f.write("\n")
                            f.write("-"*120)
                            f.write("\n")
                            f.write(str(d[i+1]))
                            f.write("\n")
                            if (choose_list[0] == "001") or (choose_list[0]== "002") or (choose_list[0]=="003")or (choose_list[0] == "004")or (choose_list[0]== "008"):
                                for i in range (len(tran_type_list)):
                                    if ((int(choose_list[0][2]) == tran_type_list[i]) and (tran_status_list[i]== "Failed") and (tran_year_list[i] == choose_list[1]) and (tran_month_list[i] == choose_list[2])and (tran_credit_card_list[i][2] == str(bank_list[0]))):
                                        count+=(tran_amount_list[i])
                                        count_customer+=1
                                f.write("\n")
                                f.write(" "*90)
                                f.write("Total amount = RM ")
                                f.write(str(count))
                                f.write("\n")
                                f.write(" "*90)
                                f.write("Total of Customer = ")
                                f.write(str(count_customer))
                                f.write("\n")
                                f.write("-"*120)
                            count=0
                            count_customer =0
                            if (choose_list[0] == "007") or (choose_list[0] == "009") or (choose_list[0] == "006"):
                                for i in range (len(tran_type_list)):
                                    if ((int(choose_list[0][2]) == tran_type_list[i]) and (tran_status_list[i]== "Failed") and (tran_year_list[i] == choose_list[1]) and (tran_month_list[i] == choose_list[2])and (tran_credit_card_list[i][2] == str(bank_list[0]))):
                                        count_customer +=1
                                f.write("\n")
                                f.write(" "*90)
                                f.write("Total of Customer = ")
                                f.write(str(count_customer))
                                f.write("\n")
                                f.write("-"*120)
                            count=0
                            count_customer=0
                    print("Please Check your file >>> File Name = ",sort_filename)
                    print("Location : Documents // this File \n\n.....................Thank You.....................")
                    enter=input("Press Enter to continue :)")
                    group_n.clear()
                    tran_error_sort_option_display(tran_file,column)
                elif (option =="2")or (option == "3"):
                    print("...................Processing...................")
                    print(" .................Please Wait..................")
                    from docx import Document
                    from docx.shared import Inches
                    word_sort_filename= sort_filename + ".docx"
                    document = Document()
                    document.add_heading("                        [ABC Company]         ",0).bold = True
                    document.add_paragraph("Report   : Transactions Error Sorted Report ").bold=True
                    document.add_paragraph("Address  :  122,Jalan Ikan , 20111 Kuala Kangsa , Malaysia")
                    document.add_paragraph("Reported by : Alex Pan Wan")
                    document.add_paragraph("_________________________________________________________________________________________________________")
                    document.add_paragraph(str(datetime.date.today()))
                    document.add_paragraph(str(time.strftime("%H:%M:%S")))
                    t = document.add_table(sort_3.shape[0]+1, sort_3.shape[1])
                    for j in range(sort_3.shape[-1]):
                        t.cell(0,j).text = sort_3.columns[j]
                    for i in range (sort_3.shape[0]):
                        for j in range(sort_3.shape[-1]):
                            t.cell(i+1,j).text = str (sort_3.values[i,j])
                    document.save(str(word_sort_filename))
                    if (option == "2"):
                        print("Please Check your file >>> File Name = ",word_sort_filename)
                        enter=input("Press Enter to continue :)")
                        tran_error_sort_option_display(tran_file,column)
                    elif (option == "3"):
                        from docx2pdf import convert
                        pdf_sort_filename= sort_filename + ".pdf"
                        convert(word_sort_filename,pdf_sort_filename)
                        os.remove(word_sort_filename)
                        print("Please Check your file >>> File Name = ",pdf_sort_filename)
                        print("Location : Documents // this File \n Thank You")
                        enter=input("Press Enter to continue :)")
                        group_n.clear()
                        tran_error_sort_option_display(tran_file,column)
                elif(option == "4"):
                    group_n.clear()
                    tran_error_sort_option_display(tran_file,column)
            if (option== "Q"):
                choose_list.clear()
                tran_error_sort_option( )
                


### B.1.2.2 - Sort Option       = 1. Transaction Type + Month + Year
###       - Transac tion Type = 10
###       - Print Sorted Credit Card Transaction error Record + Create text file // PDF file // WORD file (Error)
## This function will display all the transaction error record that user choose based on transaction type + month + year and user can print the sorted transaction record .
def tran_error_sort_option_display_10(tran_file,column):
    os.system("cls")
    global choose_list
    group_n =[]
    tran_year_list,tran_month_list ,tran_type_list,tran_merchant_id_list,tran_account_id_list,tran_status_list,tran_amount_list ,tran_credit_card_list = read_updated_transaction_file(transaction_file)
    correct_pages ,error_pages = count_page()
    tran_sort_df = pd.concat(pd.read_excel(tran_file, sheet_name=None,header = 2 ,usecols=column), ignore_index=True)
    tran_sort_df['Merchant Id number'] = tran_sort_df['Merchant Id number'].fillna(0).astype(int)
    tran_sort_df['Merchant Id number'] = (tran_sort_df['Merchant Id number'].astype(str))
    tran_sort_df['Merchant Id number'] = (tran_sort_df['Merchant Id number'].str.rjust(3, "0"))
    tran_sort_df['Merchant Id number']  = tran_sort_df['Merchant Id number'].replace("000",'Null')            
    tran_sort_df['Tran.Type'] = tran_sort_df['Tran.Type'].fillna(0).astype(int)
    tran_sort_df['Tran.Type'] = (tran_sort_df['Tran.Type'].astype(str))
    tran_sort_df['Tran.Type'] = (tran_sort_df['Tran.Type'].str.rjust(3, "0"))
    tran_sort_df['Tran.Type']  = tran_sort_df['Tran.Type'].replace("000",'Null')
    tran_sort_df['Tran.Amount/Credit Limit']= tran_sort_df['Tran.Amount/Credit Limit'].fillna(0).astype(float)
    tran_sort_df['Tran.Amount/Credit Limit'] = ('RM' + tran_sort_df['Tran.Amount/Credit Limit'].astype(str))
    tran_sort_df['Date (dd/mm/yyy)'] = pd.to_datetime(tran_sort_df['Date (dd/mm/yyy)'])
    tran_sort_df['Date (dd/mm/yyy)'] = pd.to_datetime(tran_sort_df['Date (dd/mm/yyy)'])
    tran_sort_df.fillna("Null", inplace=True)
    sort1=tran_sort_df[tran_sort_df["Status"]== "Failed" ]
    if (bank_list[0] == "1"):
        sort1 =sort1[sort1['Credit Card Number'].str.startswith("001", na = False)]
    elif(bank_list[0] == "2"):
        sort1 =sort1[sort1['Credit Card Number'].str.startswith("002", na = False)]
    elif (bank_list [0] == "3"):
        sort1 =sort1[sort1['Credit Card Number'].str.startswith("003", na = False)]
    sort_2=sort1[sort1['Date (dd/mm/yyy)'].dt.year == (choose_list[1])]
    sort_3=sort_2[sort_2['Date (dd/mm/yyy)'].dt.month == (choose_list[2])]
    b1 = (len(sort_3)/20)
    b2 = (len(sort_3)//20)
    if b1>b2:
        sort_pages=b2+1
    elif b1 == b2 :
        sort_pages=b2
    elif b2 == 0 or b1 == 0 :
        sort_pages = 1
    if (len(sort_3)==0):
        print("\n")
        print("                                           ABC Company                                         ", datetime.date.today())
        print("-"*120)
        print("Report   : Transactions Error Sorted Report ")
        print("Address  :  122,Jalan Ikan  , 20111 Kuala Kangsa , Malaysia")
        print("Year     :",choose_list[1])
        print("Month    :",choose_list[2],"   "*21,"[ Page = 1 ] ")
        print("-"*120)
        print("                    [Empty Record]            ")
        print(" "*80,"(Q)uit")
        print("-"*120)
        option=input("Option = ").upper()
        while (option != "Q"):
            print ("\t[Please enter again.]")
            option=input("Option = ").upper()
        if option == "Q" :
            choose_list.clear()
            tran_error_sort_option( )

    else:
        n = (len(sort_3)//sort_pages)
        group_n.append(0)
        if (isinstance(n,int)) == True :
            for i in range (sort_pages):
                group_n.append(n)
        else:
            for i in range (sort_pages):
                if (sort_pages - i) != 1:
                    n= (len(sort_3)//sort_pages)
                    group_n.append(n)
                else:
                    special_n=(len(sort_3)-(i*n))
                    group_n.append(special_n)
        groups = np.cumsum(group_n)
        d = {}
        for i in range(sort_pages):
            print("\n")
            print("                                           ABC Company                                         ", datetime.date.today())
            print("-"*120)
            print("Report   : Transactions  Error Sorted Report ")
            print("Address  :  122,Jalan Ikan  , 20111 Kuala Kangsa , Malaysia")
            print("Year     :",choose_list[1])
            print("Month    :",choose_list[2],"   "*21,"[ Page",sort_pages,"of",i+1,"]")
            print("-"*120)
            print("\n")
            d[i+1] = sort_3.iloc[groups[i]:groups[i+1]]
            print(d[i+1])
            print("\n")
            print("\n","  "*60,"Total record =",(len(sort_3)),"\n")
            print("-"*120)
            print(" "*80,"(N)ext  (P)rint (Q)uit")
            print("-"*120)
            option=input("Option = ").upper()
            while (option!="N") and (option != "Q") and (option != "P"):
                print ("\t[Please enter again.]")
                option=input("Option = ")
            if (option== "N"):
                if ((sort_pages - i) != 1)and (sort_pages != 1) :
                    os.system("cls")
                    continue
                if (sort_pages == 1 ):
                    tran_error_sort_option_display(tran_file,column)
                else:
                    tran_error_sort_option_display(tran_file,column)
            if (option== "P"):
                with open ("display.txt" , "r") as f :
                    for line in f.readlines()[25:33]:
                        print(line)
                option=input("Option = ")
                option_list=["1","2","3","4"]
                while (option not in option_list):
                    option=input("Option = ")
                    print ("\t[Please enter again.]")
                sort_filename="Transaction Error Sorted File_"+""+str(time.strftime("%Y%m%d-%H%M%S"))
                if (option =="1"):
                    print("...................Processing...................")
                    print(" .................Please Wait..................")
                    sort_filename= sort_filename + ".txt"
                    with open (sort_filename,'w') as f :
                        for i in range(sort_pages):
                            f.write("\n\n")
                            f.write("                                           ABC Company                                         ")
                            f.write(str(datetime.date.today()))
                            f.write(str(time.strftime("%H:%M:%S")))
                            f.write("\n")
                            f.write("-"*120)
                            f.write("\n")
                            f.write("Report           : Transactions  Error Sorted Report \n")
                            f.write("Transaction Type :")
                            f.write(str(choose_list[0]))
                            f.write("\n")
                            f.write("Address  :  122,Jalan Ikan , 20111 Kuala Kangsa , Malaysia\n")
                            f.write("Year     : ")
                            f.write(str(choose_list[1]))
                            f.write("\n")
                            f.write("Month    : ")
                            f.write(str(choose_list[2]))
                            f.write("   "*21)
                            f.write("Page ")
                            f.write(str(sort_pages))
                            f.write(" of ")
                            f.write(str(i+1))
                            f.write("\n")
                            f.write("-"*120)
                            f.write("\n")
                            f.write(str(d[i+1]))
                            f.write("\n")
                            f.write("  "*60)
                            f.write("Total record =")
                            f.write(str(len(sort_3)))
                            f.write("\n")
                    print("Please Check your file >>> File Name = ",sort_filename)
                    print("Location : Documents // this File \n\n.....................Thank You.....................")
                    enter=input("Press Enter to continue :)")
                    group_n.clear()
                    tran_error_sort_option_display_10(tran_file,column)
                elif (option =="2")or (option == "3"):
                    print("...................Processing...................")
                    print(" .................Please Wait..................")
                    from docx import Document
                    from docx.shared import Inches
                    word_sort_filename= sort_filename + ".docx"
                    document = Document()
                    document.add_heading("                        [ABC Company]         ",0).bold = True
                    document.add_paragraph("Report   : Transactions Error Sorted Report ").bold=True
                    document.add_paragraph("Address  :  122,Jalan Ikan , 20111 Kuala Kangsa , Malaysia")
                    document.add_paragraph("Reported by : Alex Pan Wan")
                    document.add_paragraph("_________________________________________________________________________________________________________")
                    document.add_paragraph(str(datetime.date.today()))
                    document.add_paragraph(str(time.strftime("%H:%M:%S")))
                    t = document.add_table(sort_3.shape[0]+1, sort_3.shape[1])
                    for j in range(sort_3.shape[-1]):
                        t.cell(0,j).text = sort_3.columns[j]
                    for i in range (sort_3.shape[0]):
                        for j in range(sort_3.shape[-1]):
                            t.cell(i+1,j).text = str (sort_3.values[i,j])
                    document.save(str(word_sort_filename))
                    if (option == "2"):
                        print("Please Check your file >>> File Name = ",word_sort_filename)
                        enter=input("Press Enter to continue :)")
                        group_n.clear()
                        tran_error_sort_option_display_10(tran_file,column)
                    elif (option == "3"):
                        from docx2pdf import convert
                        pdf_sort_filename= sort_filename + ".pdf"
                        convert(word_sort_filename,pdf_sort_filename)
                        os.remove(word_sort_filename)
                        print("Please Check your file >>> File Name = ",pdf_sort_filename)
                        print("Location : Documents // this File \n Thank You")
                        enter=input("Press Enter to continue :)")
                        group_n.clear()
                        tran_error_sort_option_display_10(tran_file,column)
                elif(option == "4"):
                    group_n.clear()
                    tran_error_sort_option_display(tran_file,column)
            if (option== "Q"):
                choose_list.clear()
                tran_error_sort_option( )



### B.2 - Sort Option   = 2. Customer Account (eg:00X-XXXX) , Month , Year
#   This function allows user to insert the customer Account ID + Month + Year 
def transaction_error_sort_option1():
    os.system("cls")
    correct_pages ,error_pages = count_page()
    tran_year_list,tran_month_list ,tran_type_list,tran_merchant_id_list,tran_account_id_list,tran_status_list,tran_amount_list ,tran_credit_card_list = read_updated_transaction_file(transaction_file)
    acc_id_list,acc_name_list,acc_ic_list,acc_credit_list,acc_balance_list,acc_status_list,acc_address_list = read_updated_account_file(account_file)
    group_n=[]
    global choose_list
    global global_list
    global loop
    print("-------------------------------------------------------------------")
    print("                       Insert Customer's Account Id                ")
    print("-------------------------------------------------------------------")
    print("              Account ID         Name    Balance     Status")
    for i in range (len(acc_id_list)):
        acc_id_list[i]=str(acc_id_list[i])
        if (acc_id_list[i][2] ==  bank_list[0]):
            print("          ","*",acc_id_list[i] ,"     ", acc_name_list[i],"        RM",acc_balance_list[i] ,"   ",acc_status_list[i])
    print("-------------------------------------------------------------------")
    acc_option=input("Insert Customer's Account Id (eg:00X-XXXX):")
    while ((acc_option not in acc_id_list) and (acc_option[2] == bank_list[0])):
        print ("\t[Please enter again.]")
        acc_option=input("Insert Customer's Account Id (eg:00X-XXXX):")
    choose_list.append(acc_option)
    print("-------------------------------------------------------------------")
    print("                        Choose Year")
    print("-------------------------------------------------------------------")
    space1_list = tran_year_list.copy()
    space1_list.sort()
    for i in range (len(space1_list)):
        if (space1_list [i] not in global_list):
            global_list.append(space1_list[i])
            print("                 ","*",space1_list[i])
    print("-------------------------------------------------------------------")
    while loop :
        try:
            year=input("Year (eg:20XX) = ")
            if (int(year) in tran_year_list):
                loop = False
            else:
                loop = True
                print ("\t[Please enter again.]")
        except ValueError:
            loop = True
            print ("\t[Please enter again.]")
    choose_list.append(int(year))
    global_list.clear()
    loop = True 
    print("-------------------------------------------------------------------")
    print("                        Choose Month")
    print("-------------------------------------------------------------------")
    space2_list = tran_month_list.copy()
    space2_list.sort()
    for i in range (len(space2_list)):
        if (space2_list [i] not in global_list):
            global_list.append(space2_list[i])
            print("                 *",space2_list[i],"=",calendar.month_name[space2_list[i]])
    print("-------------------------------------------------------------------")
    while loop :
        try:
            month=input("Month (eg:1) = ")
            if (int(month) in tran_month_list):
                loop = False
            else:
                loop = True
                print ("\t[Please enter again.]")
        except ValueError:
            loop = True
            print ("\t[Please enter again.]")
    loop = True 
    choose_list.append(int(month))
    global_list.clear()
    cols=[0,2,3,4,5,6,11,12]
    transaction_error_sort_option1_display(transaction_file,cols)



### B.2.1 - Print the Credit Card Transaction Record in TEXT file // WORD file // PDF file (No Error)
###       - Based on the customer's Account ID , Month , Year that user choose.
def  transaction_error_sort_option1_display(tran_file,column):
    global choose_list
    group_n =[]
    acc_id_list,acc_name_list,acc_ic_list,acc_credit_list,acc_balance_list,acc_status_list,acc_address_list = read_updated_account_file(account_file)
    tran_year_list,tran_month_list ,tran_type_list,tran_merchant_id_list,tran_account_id_list,tran_status_list,tran_amount_list ,tran_credit_card_list = read_updated_transaction_file(transaction_file)
    correct_pages ,error_pages = count_page()
    position = acc_id_list.index(choose_list[0])
    tran_sort_df = pd.concat(pd.read_excel(tran_file, sheet_name=None,header = 2 ,usecols=column), ignore_index=True)
    tran_sort_df['Merchant Id number'] = tran_sort_df['Merchant Id number'].fillna(0).astype(int)
    tran_sort_df['Merchant Id number'] = (tran_sort_df['Merchant Id number'].astype(str))
    tran_sort_df['Merchant Id number'] = (tran_sort_df['Merchant Id number'].str.rjust(3, "0"))
    tran_sort_df['Merchant Id number']  = tran_sort_df['Merchant Id number'].replace("000",'Null')            
    tran_sort_df['Tran.Type'] = tran_sort_df['Tran.Type'].fillna(0).astype(int)
    tran_sort_df['Tran.Type'] = (tran_sort_df['Tran.Type'].astype(str))
    tran_sort_df['Tran.Type'] = (tran_sort_df['Tran.Type'].str.rjust(3, "0"))
    tran_sort_df['Tran.Type']  = tran_sort_df['Tran.Type'].replace("000",'Null')
    tran_sort_df['Tran.Amount/Credit Limit']= tran_sort_df['Tran.Amount/Credit Limit'].fillna(0).astype(float)
    tran_sort_df['Tran.Amount/Credit Limit'] = ('RM' + tran_sort_df['Tran.Amount/Credit Limit'].astype(str))
    tran_sort_df['Date (dd/mm/yyy)'] = pd.to_datetime(tran_sort_df['Date (dd/mm/yyy)'])
    tran_sort_df.fillna("Null", inplace=True)
    sort1=tran_sort_df[tran_sort_df["Status"]== "Failed" ]
    sort_1=sort1[sort1['Credit Card Number']==(choose_list[0])]
    sort_2=sort_1[sort_1['Date (dd/mm/yyy)'].dt.year == (choose_list[1])]
    sort_3=sort_2[sort_2['Date (dd/mm/yyy)'].dt.month == (choose_list[2])]
    b1 = (len(sort_3)/20)
    b2 = (len(sort_3)//20)
    if b1>b2:
        sort_pages=b2+1
    elif b1 == b2 :
        sort_pages=b2
    elif b2 == 0 or b1 == 0 :
        sort_pages = 1
    if (len(sort_3)==0):
        print("\n")
        print("                                           ABC Company                                         ", datetime.date.today())
        print("-"*120)
        print("Report   : Transactions Error Sorted Report ")
        print("Customer Name   :",acc_name_list[position])
        print("Customer IC     :",acc_ic_list[position])
        print("Customer Add    :",acc_address_list[position])
        print("Customer Balance:",acc_balance_list[position])
        print("Address  :  122,Jalan Ikan  , 20111 Kuala Kangsa , Malaysia")
        print("Year     :",choose_list[1])
        print("Month    :",choose_list[2],"   "*21,"[ Page = 1 ]")
        print("-"*120)
        print("                    [Empty Record]            ")
        print(" "*80,"(Q)uit")
        print("-"*120)
        option=input("Option = ").upper()
        while (option != "Q"):
            print ("\t[Please enter again.]")
            option=input("Option = ").upper()
        if option == "Q" :
            choose_list.clear()
            transaction_error_report(error_pages,transaction_file)
    else:
        n = (len(sort_3)//sort_pages)
        group_n.append(0)
        if (isinstance(n,int)) == True :
            for i in range (sort_pages):
                group_n.append(n)
        else:
            for i in range (sort_pages):
                if (sort_pages - i) != 1:
                    n= (len(sort_3)//sort_pages)
                    group_n.append(n)
                else:
                    special_n=(len(sort_3)-(i*n))
                    group_n.append(special_n)
        groups = np.cumsum(group_n)
        d = {}
        for i in range(sort_pages):
            print("\n")
            print("                                           ABC Company                                         ", datetime.date.today())
            print("-"*120)
            print("Report          : Transactions Error Sorted Report ")
            print("Customer ID     :",choose_list[0])
            print("Customer Name   :",acc_name_list[position])
            print("Customer IC     :",acc_ic_list[position])
            print("Customer Add    :",acc_address_list[position])
            print("Customer Balance:",acc_balance_list[position])
            print("Address         :  122,Jalan Ikan  , 20111 Kuala Kangsa , Malaysia")
            print("Year            :",choose_list[1])
            print("Month           :",choose_list[2],"   "*21,"[ Page",sort_pages,"of",i+1,"]")
            print("-"*120)
            print("\n")
            d[i+1] = sort_3.iloc[groups[i]:groups[i+1]]
            print(d[i+1])
            count=0
            count_customer =0
            print("\n")
            print("  "*60,"Total record =",(len(sort_3)),"\n")
            print("-"*120)
            print(" "*80,"(N)ext   (P)rint    (Q)uit")
            print("-"*120)
            option=input("Option = ").upper()
            while (option!="N") and (option != "Q") and (option != "P"):
                print ("\t[Please enter again.]")
                option=input("Option = ")
            if (option== "N"):
                if ((sort_pages - i) != 1)and (sort_pages != 1) :
                    os.system("cls")
                    continue
                if (sort_pages == 1 ):
                    transaction_error_sort_option1_display(tran_file,column)
                else:
                    transaction_error_sort_option1_display(tran_file,column)
            if (option== "P"):
                with open ("display.txt" , "r") as f :
                    for line in f.readlines()[25:33]:
                        print(line)
                option=input("Option = ")
                option_list=["1","2","3","4"]
                while (option not in option_list):
                    option=input("Option = ")
                    print ("\t[Please enter again.]")
                sort_filename="Transaction Error Sorted File_"+""+str(time.strftime("%Y%m%d-%H%M%S"))
                if (option =="1"):
                    print("...................Processing...................")
                    print(" .................Please Wait..................")
                    sort_filename= sort_filename + ".txt"
                    with open (sort_filename,'w') as f :
                        for i in range(sort_pages):
                            f.write("\n\n")
                            f.write("                                           ABC Company                                         ")
                            f.write(str(datetime.date.today()))
                            f.write(str(time.strftime("%H:%M:%S")))
                            f.write("\n")
                            f.write("-"*120)
                            f.write("\n")
                            f.write("Report           : Transactions Error Sorted Report ")
                            f.write("\n")
                            f.write("Customer ID   : ")
                            f.write(str(choose_list[0]))
                            f.write("\n")
                            f.write("Customer Name  : ")
                            f.write(str(acc_name_list[position]))
                            f.write("\n")
                            f.write("Customer IC   : ")
                            f.write(str(acc_ic_list[position]))
                            f.write("\n")
                            f.write("Customer Add   : ")
                            f.write(str(acc_address_list[position]))
                            f.write("\n")
                            f.write("Customer Balance : ")
                            f.write(str(acc_balance_list[position]))
                            f.write("\n")
                            f.write("Address  :  122,Jalan Ikan , 20111 Kuala Kangsa , Malaysia\n")
                            f.write("Year     : ")
                            f.write(str(choose_list[1]))
                            f.write("\n")
                            f.write("Month    : ")
                            f.write(str(choose_list[2]))
                            f.write("   "*21)
                            f.write("Page ")
                            f.write(str(sort_pages))
                            f.write(" of ")
                            f.write(str(i+1))
                            f.write("\n")
                            f.write("-"*120)
                            f.write("\n")
                            f.write(str(d[i+1]))
                            f.write("\n")
                            f.write("  "*60)
                            f.write("Total record =")
                            f.write(str(len(sort_3)))
                            f.write("\n")
                            count=0
                            count_customer=0
                    print("Please Check your file >>> File Name = ",sort_filename)
                    print("Location : Documents // this File \n Thank You")
                    enter=input("Press Enter to continue :)")
                    group_n.clear()
                    transaction_error_sort_option1_display(tran_file,column)
                elif (option =="2")or (option == "3"):
                    print("...................Processing...................")
                    print(" .................Please Wait..................")
                    from docx import Document
                    from docx.shared import Inches
                    word_sort_filename= sort_filename + ".docx"
                    document = Document()
                    document.add_heading("                        [ABC Company]         ",0).bold = True
                    document.add_paragraph("Report   : Transactions Error Sorted Report ").bold=True
                    document.add_paragraph("Address  :  122,Jalan Ikan , 20111 Kuala Kangsa , Malaysia")
                    document.add_paragraph("Reported by : Alex Pan Wan")
                    document.add_paragraph("_________________________________________________________________________________________________________")
                    document.add_paragraph(str(datetime.date.today()))
                    document.add_paragraph(str(time.strftime("%H:%M:%S")))
                    t = document.add_table(sort_3.shape[0]+1, sort_3.shape[1])
                    for j in range(sort_3.shape[-1]):
                        t.cell(0,j).text = sort_3.columns[j]
                    for i in range (sort_3.shape[0]):
                        for j in range(sort_3.shape[-1]):
                            t.cell(i+1,j).text = str (sort_3.values[i,j])
                    document.save(str(word_sort_filename))
                    if (option == "2"):
                        print("Please Check your file >>> File Name = ",word_sort_filename)
                        enter=input("Press Enter to continue :)")
                        group_n.clear()
                        transaction_error_sort_option1_display(tran_file,column)
                    elif (option == "3"):
                        from docx2pdf import convert
                        pdf_sort_filename= sort_filename + ".pdf"
                        convert(word_sort_filename,pdf_sort_filename)
                        os.remove(word_sort_filename)
                        print("Please Check your file >>> File Name = ",pdf_sort_filename)
                        print("Location : Documents // this File \n Thank You")
                        enter=input("Press Enter to continue :)")
                        group_n.clear()
                        transaction_error_sort_option1_display(tran_file,column)
                elif(option == "4"):
                    group_n.clear()
                    transaction_error_sort_option1_display(tran_file,column)
            if (option== "Q"):
                choose_list.clear()
                transaction_error_report(error_pages,transaction_file)



# 4. [ Choose Customer's Account Information and Balances ] - Active // Not Active
# This function allows user to choose to view the customer's account infomation that is active or not active . 
def report1_option():
    os.system("cls")
    option_list=["1","2","3"]
    with open ("display.txt" , "r") as f :
        for line in f.readlines()[51:59]:
            print(line)
    option=input("Option = ")
    while (option not in option_list):
        print ("\t[ Please Try Again .]") 
        option=input("Option = ")
    if (option == option_list[0]):
        correct_pages ,error_pages = count_acc_page()
        acc_df ,acc_df2 ,count= account_report(correct_pages,account_file)
    elif (option == option_list[1]):
        correct_pages ,error_pages = count_acc_page()
        ac_df ,ac_df2 ,count = account_closed_report(error_pages,account_file)
    elif (option == option_list[2]):
        choose_record()


###C.1 [ Customer's Account Information and Balances  (Active)]
## This function will display all the active customer's Account information .
def account_report(c_pages,acc_file):
    os.system("cls")
    global calendar
    group_n=[]
    cols=[0,1,2,3,4,5,6,9]
    count=0
    acc_id_list,acc_name_list,acc_ic_list,acc_credit_list,acc_balance_list,acc_status_list,acc_address_list = read_updated_account_file(account_file)
    count_acc_page()
    acc_df = pd.concat(pd.read_excel(acc_file, sheet_name=None,header = 2 ,usecols=cols), ignore_index=True)
    acc_df.fillna("Null", inplace=True)
    if (bank_list[0] == "1"):
        acc_df = acc_df[acc_df['Account Id'].str.startswith("001", na = False)]
    elif(bank_list[0] == "2"):
        acc_df = acc_df[acc_df['Account Id'].str.startswith("002", na = False)]
    elif (bank_list [0] == "3"):
        acc_df = acc_df[acc_df['Account Id'].str.startswith("003", na = False)]
    acc_df['IC.No']=acc_df['IC.No'].astype(int)
    acc_df['Current PIN']=acc_df['Current PIN'].astype(int)
    acc_df['Credit Limit']=acc_df['Credit Limit'].fillna(0).astype(float)
    acc_df['Credit Limit'] = ('RM' +  acc_df['Credit Limit'].astype(str))
    acc_df['Account Balance(Amount)']=acc_df['Account Balance(Amount)'].fillna(0).astype(float)
    acc_df['Account Balance(Amount)'] = ('RM' +  acc_df['Account Balance(Amount)'].astype(str))
    acc_df2=acc_df[acc_df['Credit .C Status']== "Active" ]
    pd.set_option('display.max_rows', None)
    pd.set_option('display.max_columns', None)
    pd.set_option('display.width', None)
    pd.set_option('display.max_colwidth', None)
    if (len(acc_df2)==0):
        print("\n")
        print("                                           ABC Company                                 ", datetime.date.today() ,time.strftime("%H:%M:%S"))
        print("-"*120)
        print("Report   :Customer's Account (Active) Report ")
        print("Address  :  122,Jalan Ikan  , 20111 Kuala Kangsa , Malaysia")
        print(" "*30,"[ Page = 1]")
        print("-"*120)
        print("\n\n")
        print("                    [Empty Record]            ")
        print(" "*75,"(Q)uit")
        print("-"*120)
        option = input("Option = ")
        while (option != "Q"):
            print ("\t[Please enter again.]")
            option=input("Option = ").upper()
        if option == "Q" :
            choose_list.clear()
            report1_option()
    print(len(acc_df2))
    if (len(acc_df2)!=0):
        n= (len(acc_df2)//c_pages)
        group_n.append(0)
        if (isinstance(n,int)) == True :
            for i in range (c_pages):
                group_n.append(n)
        else:
            for i in range (c_pages):
                if (c_pages - i) != 1:
                    n= (len(acc_df2)//c_pages)
                    group_n.append(n)
                else:
                    special_n=(len(acc_df2)-(i*n))
                    group_n.append(special_n)
        groups = np.cumsum(group_n)
        d = {}
        for i in range (len(acc_status_list)):
            if (acc_status_list[i] == "Active"):
                count+=1
        for i in range(c_pages):
            print("\n")
            print("                                           ABC Company                                 ", datetime.date.today() ,time.strftime("%H:%M:%S"))
            print("-"*120)
            print("Report   :Customer's Account (Active) Report ")
            print("Address  :  122,Jalan Ikan  , 20111 Kuala Kangsa , Malaysia")
            print("   "*30,"[ Page",c_pages,"of",i+1,"]")
            print("-"*120)
            print("\n\n")
            d[i+1] = acc_df2.iloc[groups[i]:groups[i+1]]
            print(d[i+1])
            print("\n","   "*28,"Total record =",(count),"\n")
            print(" "*70,"(N)ext   (P)rint   (Q)uit")
            print("-"*120)
            option=input("Option = ").upper()
            option_list=["N","P","Q"]
            while option not in option_list:
                print ("\t[Please enter again.]")
                option=input("Option = ").upper()
            if(option == "N"):
                if (c_pages - i) != 1:
                    os.system("cls")
                    continue
                else:
                    account_report(c_pages,acc_file)
            elif (option == "P"):
                group_n.clear()
                acc_print_option(acc_df,acc_df2,c_pages,count)
            elif (option == "Q"):
                group_n.clear()
                report1_option()
    return acc_df,acc_df2,count


###C.1.1 Print Customer's Account Information and Balances  (Active) in TEXT file // WORD file // PDF file (No Error)
def acc_print_option(data1,data,c_pages,count):
    os.system("cls")
    correct_pages ,error_pages = count_acc_page()
    group_n=[]
    acc_id_list,acc_name_list,acc_ic_list,acc_credit_list,acc_balance_list,acc_status_list,acc_address_list = read_updated_account_file(account_file)
    with open ("display.txt" , "r") as f :
        for line in f.readlines()[25:33]:
            print(line)
    option=input("Option = ")
    filename="Customer's Account(Active) File_"+""+str(time.strftime("%Y%m%d-%H%M%S"))
    option_list=["1","2","3","4"]
    while option not in option_list:
        print ("\t[Please enter again.]")
        option=input("Option = ")
    if (option =="1"):
        filename=filename + ".txt"        
        n= (len(data)//c_pages)
        group_n.append(0)
        if (isinstance(n,int)) == True :
            for i in range (c_pages):
                group_n.append(n)
        else:
            for i in range (c_pages):
                if (c_pages - i) != 1:
                    n= (len(data)//c_pages)
                    group_n.append(n)
                else:
                    special_n=(len(data)-(i*n))
                    group_n.append(special_n)
        groups = np.cumsum(group_n)
        with open (filename,"w") as f :
            for i in range(c_pages):
                d = {}
                d[i+1]= data.iloc[groups[i]:groups[i+1]]
                f.write("\n\n\n")
                f.write("                                           ABC Company                                         ")
                f.write(str(datetime.date.today()))
                f.write(str(time.strftime("%H:%M:%S")))
                f.write("\n")
                f.write("-"*120)
                f.write("\n")
                f.write("Report   : Customer's Account (Active) Report  \n")
                f.write("Address  :  122,Jalan Ikan , 20111 Kuala Kangsa , Malaysia\n")
                f.write("Reported by  : Alex Pan Wan ")
                f.write("   "*21)
                f.write("Page ")
                f.write(str(c_pages))
                f.write(" of ")
                f.write(str(i+1))
                f.write("\n")
                f.write("-"*120)
                f.write("\n")
                f.write(str(d[i+1]))
                f.write("\n\n")
                f.write("  "*40)
                f.write("Total record = ")
                f.write(str(count))
                f.write("\n")
                f.write("-"*120)
                f.write("\n\n\n")
        print("Please Check your file >>> File Name = ",filename)
        print("Location : Documents // this File \n\n .......Thank You.....")
        enter=input("Press Enter to continue :)")
        group_n.clear()
        acc_print_option(data1,data,c_pages,count)
    elif (option =="2") or (option == "3"):
        print("...................Processing...................")
        print(" .................Please Wait..................")
        from docx import Document
        from docx.shared import Inches
        word_filename= filename + ".docx"
        document = Document()
        document.add_heading("                        [ABC Company]         ",0).bold = True
        document.add_paragraph("Report   :Customer's Account (Active) Report ").bold=True
        document.add_paragraph("Address  :  122,Jalan Ikan , 20111 Kuala Kangsa , Malaysia")
        document.add_paragraph("Reported by : Alex Pan Wan")
        document.add_paragraph("_________________________________________________________________________________________________________")
        document.add_paragraph(str(datetime.date.today()))
        document.add_paragraph(str(time.strftime("%H:%M:%S")))
        t = document.add_table(data.shape[0]+1, data.shape[1])
        for j in range(data.shape[-1]):
            t.cell(0,j).text = data.columns[j]
        for i in range (data.shape[0]):
            for j in range(data.shape[-1]):
                t.cell(i+1,j).text = str (data.values[i,j])
        document.save(str(word_filename))
        if (option == "2"):
            print("Please Check your file >>> File Name = ",word_filename)
            enter=input("Press Enter to continue :)")
            acc_print_option(data1,data,c_pages,count)
        elif (option == "3"):
            from docx2pdf import convert
            pdf_filename=filename + ".pdf"
            convert(word_filename,pdf_filename)
            os.remove(word_filename)
            print("Please Check your file >>> File Name = ",pdf_filename)
            print("Location : Documents // this File \n\n............ Thank You.............")
            enter=input("Press Enter to continue :)")
            acc_print_option(data1,data,c_pages,count)
    elif (option == "4"):
        account_report(correct_pages,account_file)


###C.2 [ Customer's Account Information and Balances  (Closed)]
## This function will display all the active customer's Account information .      
def account_closed_report(e_pages,acc_file):
    os.system("cls")
    global calendar
    group_n=[]
    cols=[0,1,2,3,4,5,6,9]
    count=0
    acc_id_list,acc_name_list,acc_ic_list,acc_credit_list,acc_balance_list,acc_status_list,acc_address_list = read_updated_account_file(account_file)
    count_acc_page()
    ac_df = pd.concat(pd.read_excel(acc_file, sheet_name=None,header = 2 ,usecols=cols), ignore_index=True)
    ac_df.fillna("Null", inplace=True)
    if (bank_list[0] == "1"):
        ac_df = ac_df[ac_df['Account Id'].str.startswith("001", na = False)]
    elif(bank_list[0] == "2"):
        ac_df = ac_df[ac_df['Account Id'].str.startswith("002", na = False)]
    elif (bank_list [0] == "3"):
        ac_df = ac_df[ac_df['Account Id'].str.startswith("003", na = False)]
    ac_df['Current PIN']=ac_df['Current PIN'].astype(int)
    ac_df['IC.No']=ac_df['IC.No'].astype(int)
    ac_df['Credit Limit']=ac_df['Credit Limit'].fillna(0).astype(float)
    ac_df['Credit Limit'] = ('RM' +  ac_df['Credit Limit'].astype(str))
    ac_df['Account Balance(Amount)']=ac_df['Account Balance(Amount)'].fillna(0).astype(float)
    ac_df['Account Balance(Amount)'] = ('RM' +  ac_df['Account Balance(Amount)'].astype(str))
    ac_df2=ac_df[ac_df['Credit .C Status']== "Closed" ]
    pd.set_option('display.max_rows', None)
    pd.set_option('display.max_columns', None)
    pd.set_option('display.width', None)
    pd.set_option('display.max_colwidth', None)
    if (len(ac_df2)==0):
        print("\n")
        print("                                           ABC Company                                 ", datetime.date.today() ,time.strftime("%H:%M:%S"))
        print("-"*120)
        print("Report   :Customer's Account (Closed) Report ")
        print("Address  :  122,Jalan Ikan  , 20111 Kuala Kangsa , Malaysia")
        print("  "*30,"[ Page = 1]")
        print("-"*120)
        print("\n\n")
        print("                    [Empty Record]            ")
        print(" "*75,"(Q)uit")
        print("-"*120)
        option = input("Option = ")
        while (option != "Q"):
            print ("\t[Please enter again.]")
            option=input("Option = ").upper()
        if option == "Q" :
            choose_list.clear()
            report1_option()
    else:
        n= (len(ac_df2)//e_pages)
        group_n.append(0)
        if (isinstance(n,int)) == True :
            for i in range (e_pages):
                group_n.append(n)
        else:
            for i in range (e_pages):
                if (e_pages - i) != 1:
                    n= (len(ac_df2)//e_pages)
                    group_n.append(n)
                else:
                    special_n=(len(ac_df2)-(i*n))
                    group_n.append(special_n)
        groups = np.cumsum(group_n)
        d = {}
        for i in range (len(acc_status_list)):
            if (acc_status_list[i] == "Closed"):
                count+=1
        for i in range(e_pages):
            print("\n")
            print("                                           ABC Company                                 ", datetime.date.today() ,time.strftime("%H:%M:%S"))
            print("-"*120)
            print("Report   :Customer's Account (Closed) Report ")
            print("Address  :  122,Jalan Ikan  , 20111 Kuala Kangsa , Malaysia")
            print("   "*30,"[ Page",e_pages,"of",i+1,"]")
            print("-"*120)
            print("\n\n")
            d[i+1] = ac_df2.iloc[groups[i]:groups[i+1]]
            print(d[i+1])
            print("\n","   "*28,"Total record =",(count),"\n")
            print(" "*70,"(N)ext   (P)rint   (Q)uit")
            print("-"*120)
            option=input("Option = ").upper()
            option_list=["N","P","Q"]
            while option not in option_list:
                print ("\t[Please enter again.]")
                option=input("Option = ").upper()
            if(option == "N"):
                if (e_pages - i) != 1:
                    os.system("cls")
                    continue
                else:
                    account_closed_report(e_pages,acc_file)
            elif (option == "P"):
                group_n.clear()
                acc_closed_print_option(ac_df,ac_df2,e_pages,count)
            elif (option == "Q"):
                group_n.clear()
                report1_option()
    return ac_df,ac_df2,count

###C.2.1 Print customer's Account Information and Balances  (Closed) in TEXT file // WORD file // PDF file (No Error)
def acc_closed_print_option(data1,data,e_pages,count):
    os.system("cls")
    correct_pages ,error_pages = count_acc_page()
    group_n=[]
    acc_id_list,acc_name_list,acc_ic_list,acc_credit_list,acc_balance_list,acc_status_list,acc_address_list = read_updated_account_file(account_file)
    with open ("display.txt" , "r") as f :
        for line in f.readlines()[25:33]:
            print(line)
    option=input("Option = ")
    filename="Customer's Account(Closed) File_"+""+str(time.strftime("%Y%m%d-%H%M%S"))
    option_list=["1","2","3","4"]
    while option not in option_list:
        print ("\t[Please enter again.]")
        option=input("Option = ")
    if (option =="1"):
        filename=filename + ".txt"        
        n= (len(data)//e_pages)
        group_n.append(0)
        if (isinstance(n,int)) == True :
            for i in range (e_pages):
                group_n.append(n)
        else:
            for i in range (e_pages):
                if (e_pages - i) != 1:
                    n= (len(data)//e_pages)
                    group_n.append(n)
                else:
                    special_n=(len(data)-(i*n))
                    group_n.append(special_n)
        groups = np.cumsum(group_n)
        with open (filename,"w") as f :
            for i in range(e_pages):
                d = {}
                d[i+1]= data.iloc[groups[i]:groups[i+1]]
                f.write("\n\n\n")
                f.write("                                           ABC Company                                         ")
                f.write(str(datetime.date.today()))
                f.write(str(time.strftime("%H:%M:%S")))
                f.write("\n")
                f.write("-"*120)
                f.write("\n")
                f.write("Report   : Customer's Account (Closed) Report  \n")
                f.write("Address  :  122,Jalan Ikan , 20111 Kuala Kangsa , Malaysia\n")
                f.write("Reported by  : Alex Pan Wan ")
                f.write("   "*21)
                f.write("Page ")
                f.write(str(e_pages))
                f.write(" of ")
                f.write(str(i+1))
                f.write("\n")
                f.write("-"*120)
                f.write("\n")
                f.write(str(d[i+1]))
                f.write("\n\n")
                f.write("  "*40)
                f.write("Total record = ")
                f.write(str(count))
                f.write("\n")
                f.write("-"*120)
                f.write("\n\n\n")
        print("Please Check your file >>> File Name = ",filename)
        print("Location : Documents // this File \n\n .......Thank You.....")
        enter=input("Press Enter to continue :)")
        group_n.clear()
        acc_closed_print_option(data1,data,e_pages,count)
    elif (option =="2") or (option == "3"):
        print("...................Processing...................")
        print(" .................Please Wait..................")
        from docx import Document
        from docx.shared import Inches
        word_filename= filename + ".docx"
        document = Document()
        document.add_heading("                        [ABC Company]         ",0).bold = True
        document.add_paragraph("Report   :Customer's Account (Closed) Report ").bold=True
        document.add_paragraph("Address  :  122,Jalan Ikan , 20111 Kuala Kangsa , Malaysia")
        document.add_paragraph("Reported by : Alex Pan Wan")
        document.add_paragraph("_________________________________________________________________________________________________________")
        document.add_paragraph(str(datetime.date.today()))
        document.add_paragraph(str(time.strftime("%H:%M:%S")))
        t = document.add_table(data.shape[0]+1, data.shape[1])
        for j in range(data.shape[-1]):
            t.cell(0,j).text = data.columns[j]
        for i in range (data.shape[0]):
            for j in range(data.shape[-1]):
                t.cell(i+1,j).text = str (data.values[i,j])
        document.save(str(word_filename))
        if (option == "2"):
            print("Please Check your file >>> File Name = ",word_filename)
            enter=input("Press Enter to continue :)")
            acc_closed_print_option(data1,data,e_pages,count)
        elif (option == "3"):
            from docx2pdf import convert
            pdf_filename=filename + ".pdf"
            convert(word_filename,pdf_filename)
            os.remove(word_filename)
            print("Please Check your file >>> File Name = ",pdf_filename)
            print("Location : Documents // this File \n\n............ Thank You.............")
            enter=input("Press Enter to continue :)")
            acc_closed_print_option(data1,data,e_pages,count)
    elif (option == "4"):
        account_closed_report(error_pages,account_file)
        
def main():
    global staff_name, file1, file2, file3, sheet, pd, load_workbook, option
    import pandas as pd
    from openpyxl import  load_workbook
    import os
    option, staff_name = mainmenu_1()
    file2 = 'GA2_Accounts.xlsx'
    file3 = 'Processed Transaction.xlsx'
    loop = True
    while loop:
        if   option == '1':
            while True:
                file1 = input("Please enter the excel file you want to process  : ")
                if file1.upper() == 'Q':
                    option, loop = mainmenu_2()
                    break
                else:
                    sheet = input("Please enter the excel sheet you want to process : ")
                    if sheet.upper() == 'Q':
                        option, loop = mainmenu_2()
                        break
                    else:
                        try:
                            excel_data_df  = pd.read_excel(file1, sheet_name = sheet, header = 2)
                        except:
                            print('Please enter a valid file/sheet')
                        else:
                            option1()
                            
        elif option == '2':
            global transaction_file
            global account_file
            transaction_file = 'Processed Transaction.xlsx'
            account_file = file2
            tran_df= pd.concat(pd.read_excel(transaction_file, sheet_name=None,header = 2), ignore_index=True)
            if (len(tran_df) == 0):
                print("There is no record in this file. Please process the transactions record first (option 1).")
                option=input("Please Press Enter to continue")
                os.system("cls")
                option, loop = mainmenu_2()
            if (len(tran_df) != 0):
                bank_option() 
        else:
            print('')
            option, staff_name = mainmenu_1()

main()






